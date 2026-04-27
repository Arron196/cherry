from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

import win32com.client as win32
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE_DOC = ROOT / "论文.docx"
TEMPLATE_DOC = ROOT / "论文——模板编辑.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版.docx"


CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百零]+章\s+(.+)$")
H2_RE = re.compile(r"^(\d+\.\d+)\s*(.+)$")
H3_RE = re.compile(r"^(\d+\.\d+\.\d+)\s*(.+)$")


def split_label_value(text: str) -> tuple[str, str]:
    parts = re.split(r"[：:]", text, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return text.strip(), ""


def iter_block_items(doc: DocumentObject):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def clear_document_body(doc: DocumentObject) -> None:
    body = doc.element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def ensure_rfonts(target) -> OxmlElement:
    r_pr = target._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    return r_fonts


def set_style_rfonts(style, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    r_fonts = ensure_rfonts(style)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def apply_run_font(run, size_pt: float, bold: bool = False, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_fonts = ensure_rfonts(run)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: DocumentObject) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    set_style_rfonts(normal, east_asia="宋体", ascii_font="Times New Roman")
    normal_format = normal.paragraph_format
    normal_format.line_spacing = 1.5
    normal_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)

    heading_specs = {
        "论文一级标题": (15, 0),
        "论文二级标题": (14, 1),
        "论文三级标题": (12, 2),
    }
    for style_name, (size, outline_level) in heading_specs.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        set_style_rfonts(style, east_asia="宋体", ascii_font="Times New Roman")
        fmt = style.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.line_spacing = 1
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        fmt.first_line_indent = Cm(0)
        fmt.left_indent = Cm(0)
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pr = style._element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), str(outline_level))

    # 目录内页标题
    if "TOC Heading" in doc.styles:
        toc_heading = doc.styles["TOC Heading"]
        toc_heading.font.name = "Times New Roman"
        toc_heading.font.size = Pt(16)
        toc_heading.font.bold = True
        toc_heading.font.color.rgb = RGBColor(0, 0, 0)
        set_style_rfonts(toc_heading, east_asia="宋体", ascii_font="Times New Roman")
        toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for toc_style_name in ("toc 1", "toc 2", "toc 3"):
        try:
            toc_style = doc.styles[toc_style_name]
        except KeyError:
            continue
        toc_style.font.name = "Times New Roman"
        toc_style.font.size = Pt(12)
        toc_style.font.bold = False
        toc_style.font.color.rgb = RGBColor(0, 0, 0)
        set_style_rfonts(toc_style, east_asia="宋体", ascii_font="Times New Roman")

    for link_style_name in ("Hyperlink", "FollowedHyperlink"):
        try:
            link_style = doc.styles[link_style_name]
        except KeyError:
            continue
        link_style.font.name = "Times New Roman"
        link_style.font.color.rgb = RGBColor(0, 0, 0)
        link_style.font.underline = False
        set_style_rfonts(link_style, east_asia="宋体", ascii_font="Times New Roman")


def apply_page_setup(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = False


def add_text_paragraph(
    doc: DocumentObject,
    text: str = "",
    *,
    style: str | None = None,
    size_pt: float = 12,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent_cm: float | None = None,
    left_indent_cm: float | None = None,
    line_spacing: float | None = 1.5,
    line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> Paragraph:
    p = doc.add_paragraph()
    if style:
        p.style = style
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = line_rule
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    else:
        pf.first_line_indent = None
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    else:
        pf.left_indent = None
    run = p.add_run(text)
    apply_run_font(run, size_pt=size_pt, bold=bold)
    return p


def add_toc_placeholder(doc: DocumentObject) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("[[TOC]]")
    apply_run_font(run, size_pt=12, bold=False)


def style_body_paragraph(paragraph: Paragraph, size_pt: float = 12) -> None:
    paragraph.style = doc.styles["Normal"]  # type: ignore[name-defined]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.85)
    fmt.left_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        apply_run_font(run, size_pt=size_pt, bold=bool(run.bold))


def style_reference_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = doc.styles["Normal"]  # type: ignore[name-defined]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0.74)
    fmt.first_line_indent = Cm(-0.74)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        apply_run_font(run, size_pt=10.5, bold=False)


def set_table_borders(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ("left", "right", "insideH", "insideV", "top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        if edge in {"left", "right", "insideV"}:
            el.set(qn("w:val"), "nil")
        elif edge == "insideH":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)

    # 三线表：顶部粗线、表头下细线、底部粗线，不加左右竖线。
    for r_idx, row in enumerate(table.rows):
        is_first = r_idx == 0
        is_last = r_idx == len(table.rows) - 1
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            else:
                for child in list(tc_borders):
                    tc_borders.remove(child)

            for edge in ("left", "right", "top", "bottom", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                tc_borders.append(el)

            if is_first:
                top = tc_borders.find(qn("w:top"))
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "12")
                top.set(qn("w:color"), "000000")

                bottom = tc_borders.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:color"), "000000")

            if is_last:
                bottom = tc_borders.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "12")
                bottom.set(qn("w:color"), "000000")


def style_table(table: Table) -> None:
    table.alignment = 1  # center
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt = p.paragraph_format
                fmt.first_line_indent = None
                fmt.left_indent = Cm(0)
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1.15
                fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if not p.runs:
                    p.add_run("")
                for run in p.runs:
                    apply_run_font(run, size_pt=10.5, bold=bool(run.bold))


def build_output_doc() -> Path:
    shutil.copy2(TEMPLATE_DOC, OUTPUT_DOC)

    global doc
    doc = Document(OUTPUT_DOC)
    src = Document(SOURCE_DOC)

    clear_document_body(doc)
    configure_styles(doc)

    title = src.paragraphs[0].text.strip()
    metadata_lines = [p.text.strip() for p in src.paragraphs[1:8] if p.text.strip()]
    meta_map = dict(split_label_value(line) for line in metadata_lines)

    paragraphs = [p.text.strip() for p in src.paragraphs]
    i_abs = paragraphs.index("摘要")
    i_kw_cn = next(i for i in range(i_abs + 1, len(paragraphs)) if paragraphs[i].startswith("关键词"))
    i_abs_en = paragraphs.index("Abstract")
    i_kw_en = next(i for i in range(i_abs_en + 1, len(paragraphs)) if paragraphs[i].startswith("Keywords"))

    cn_abstract = "\n".join(p for p in paragraphs[i_abs + 1 : i_kw_cn] if p)
    cn_keywords = split_label_value(paragraphs[i_kw_cn])[1]
    en_abstract = "\n".join(p for p in paragraphs[i_abs_en + 1 : i_kw_en] if p)
    en_keywords = split_label_value(paragraphs[i_kw_en])[1]

    student_name = meta_map.get("学生姓名", "")
    major = meta_map.get("专业", "")
    advisor = meta_map.get("指导教师", "")

    english_title = "Design and Implementation of a High-Quality Cherry Supply Chain Traceability System Based on STM32 and Blockchain"
    english_major_author = f"Electronic Science and Technology {student_name}".strip()
    english_supervisor = f"Supervisor: {advisor}".strip()

    # 第 1 节：封面 + 目录（不编页码）
    apply_page_setup(doc.sections[0])
    add_text_paragraph(doc, "", size_pt=12, line_spacing=None, space_after_pt=28)
    add_text_paragraph(
        doc,
        title,
        size_pt=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=18,
    )
    for line in metadata_lines:
        add_text_paragraph(
            doc,
            line,
            size_pt=12,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=None,
            space_after_pt=10,
        )
    add_text_paragraph(doc, "", size_pt=12, line_spacing=None, space_after_pt=0)
    doc.add_page_break()

    add_text_paragraph(
        doc,
        "目  录",
        style="TOC Heading",
        size_pt=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=12,
    )
    add_toc_placeholder(doc)

    # 第 2 节：摘要 + 正文（阿拉伯数字页码）
    doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(doc.sections[1])
    doc.sections[1].header.is_linked_to_previous = False
    doc.sections[1].footer.is_linked_to_previous = False

    add_text_paragraph(
        doc,
        title,
        size_pt=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=10,
    )
    add_text_paragraph(
        doc,
        f"{major} {student_name}".strip(),
        size_pt=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=6,
    )
    add_text_paragraph(
        doc,
        f"导师：{advisor}".strip("："),
        size_pt=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=12,
    )

    add_text_paragraph(
        doc,
        "摘要",
        style="论文一级标题",
        size_pt=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0,
        left_indent_cm=0,
        line_spacing=None,
        space_before_pt=0,
        space_after_pt=6,
    )
    p = add_text_paragraph(doc, cn_abstract, size_pt=12, bold=False)
    style_body_paragraph(p)

    add_text_paragraph(
        doc,
        "关键词",
        style="论文一级标题",
        size_pt=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=0,
        left_indent_cm=0,
        line_spacing=None,
        space_before_pt=6,
        space_after_pt=4,
    )
    p = add_text_paragraph(doc, cn_keywords, size_pt=12, bold=False)
    style_body_paragraph(p)

    add_text_paragraph(doc, "", size_pt=12, line_spacing=None, space_after_pt=6)
    add_text_paragraph(
        doc,
        english_title,
        size_pt=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=8,
    )
    add_text_paragraph(
        doc,
        english_major_author,
        size_pt=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=6,
    )
    add_text_paragraph(
        doc,
        english_supervisor,
        size_pt=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=None,
        space_after_pt=12,
    )

    add_text_paragraph(
        doc,
        "Abstract",
        style="论文一级标题",
        size_pt=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0,
        left_indent_cm=0,
        line_spacing=None,
        space_before_pt=0,
        space_after_pt=6,
    )
    p = add_text_paragraph(doc, en_abstract, size_pt=12, bold=False)
    style_body_paragraph(p)

    add_text_paragraph(
        doc,
        "Keywords",
        style="论文一级标题",
        size_pt=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=0,
        left_indent_cm=0,
        line_spacing=None,
        space_before_pt=6,
        space_after_pt=4,
    )
    p = add_text_paragraph(doc, en_keywords, size_pt=12, bold=False)
    style_body_paragraph(p)

    doc.add_page_break()

    start_copy = False
    in_references = False
    chapter_no = 0

    for block in iter_block_items(src):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not start_copy:
                if CHAPTER_RE.match(text):
                    start_copy = True
                else:
                    continue

            if not text:
                continue

            if text == "参考文献":
                in_references = True
                p = add_text_paragraph(
                    doc,
                    text,
                    style="论文一级标题",
                    size_pt=15,
                    bold=True,
                    line_spacing=None,
                )
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
                continue

            chapter_match = CHAPTER_RE.match(text)
            h3_match = H3_RE.match(text)
            h2_match = H2_RE.match(text)

            if chapter_match and not in_references:
                chapter_no += 1
                heading_text = f"{chapter_no} {chapter_match.group(1).strip()}"
                p = add_text_paragraph(
                    doc,
                    heading_text,
                    style="论文一级标题",
                    size_pt=15,
                    bold=True,
                    line_spacing=None,
                )
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
            elif h3_match and not in_references:
                p = add_text_paragraph(
                    doc,
                    text,
                    style="论文三级标题",
                    size_pt=12,
                    bold=True,
                    line_spacing=None,
                )
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
            elif h2_match and not in_references:
                p = add_text_paragraph(
                    doc,
                    text,
                    style="论文二级标题",
                    size_pt=14,
                    bold=True,
                    line_spacing=None,
                )
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)
            else:
                p = add_text_paragraph(doc, text, size_pt=12, bold=False)
                if in_references:
                    style_reference_paragraph(p)
                else:
                    style_body_paragraph(p)
        else:
            if not start_copy:
                continue
            new_tbl = deepcopy(block._tbl)
            doc.element.body.insert(-1, new_tbl)

    # 统一表格格式
    for table in doc.tables:
        style_table(table)

    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC


def update_toc_and_pagination(path: Path) -> None:
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    wd_header_footer_primary = 1
    wd_align_center = 1
    wd_field_page = 33
    wd_collapse_end = 0

    document = None
    try:
        document = word.Documents.Open(str(path), ReadOnly=False, Visible=False)

        for index in range(1, document.Sections.Count + 1):
            sec = document.Sections(index)
            header = sec.Headers(wd_header_footer_primary)
            footer = sec.Footers(wd_header_footer_primary)
            header.Range.Text = ""
            footer.Range.Text = ""
            footer.LinkToPrevious = False
            if index == 2:
                rng = footer.Range
                rng.Text = ""
                rng.ParagraphFormat.Alignment = wd_align_center
                rng.Collapse(wd_collapse_end)
                rng.Fields.Add(rng, wd_field_page)
                rng.Font.Name = "Times New Roman"
                rng.Font.Size = 9
                footer.PageNumbers.RestartNumberingAtSection = True
                footer.PageNumbers.StartingNumber = 1

        find_range = document.Content
        find_range.Find.ClearFormatting()
        find_range.Find.Text = "[[TOC]]"
        if find_range.Find.Execute():
            toc_range = find_range
            toc_range.Text = ""
            document.TablesOfContents.Add(
                Range=toc_range,
                UseHeadingStyles=False,
                UseOutlineLevels=False,
                AddedStyles="论文一级标题,1,论文二级标题,2,论文三级标题,3",
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
                RightAlignPageNumbers=True,
                UseHyperlinks=False,
                HidePageNumbersInWeb=True,
            )

        document.Repaginate()
        for idx in range(1, document.TablesOfContents.Count + 1):
            document.TablesOfContents(idx).Update()
        document.Fields.Update()
        document.Save()
        document.Close(False)
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        try:
            word.Quit()
        except Exception:
            pass


def main() -> None:
    out = build_output_doc()
    update_toc_and_pagination(out)
    print(out)


if __name__ == "__main__":
    main()
