from __future__ import annotations

import math
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image, ImageColor, ImageDraw, ImageFilter, ImageFont


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版.docx"
DRAFT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v2_草稿.docx"
FIGURE_DIR = ROOT / "generated_figures"


def hex_color(value: str) -> tuple[int, int, int]:
    return ImageColor.getrgb(value)


PALETTE = {
    "bg_top": hex_color("#F7FBFF"),
    "bg_bottom": hex_color("#EEF6FF"),
    "text": hex_color("#183B56"),
    "muted": hex_color("#4A647A"),
    "line": hex_color("#7AA7C7"),
    "blue": hex_color("#2563EB"),
    "cyan": hex_color("#0891B2"),
    "teal": hex_color("#14B8A6"),
    "orange": hex_color("#F97316"),
    "green": hex_color("#16A34A"),
    "amber": hex_color("#F59E0B"),
    "red": hex_color("#DC2626"),
    "slate": hex_color("#CBD5E1"),
    "white": hex_color("#FFFFFF"),
}


@dataclass
class Box:
    x: int
    y: int
    w: int
    h: int
    fill: tuple[int, int, int]
    title: str
    subtitle: str = ""
    title_color: tuple[int, int, int] = PALETTE["text"]
    subtitle_color: tuple[int, int, int] = PALETTE["muted"]


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


FONT_TITLE = get_font(56, bold=True)
FONT_H2 = get_font(40, bold=True)
FONT_BODY = get_font(28)
FONT_SMALL = get_font(24)
FONT_CODE = get_font(30, bold=True)


def lerp(a: int, b: int, t: float) -> int:
    return round(a + (b - a) * t)


def vertical_gradient(width: int, height: int, top: tuple[int, int, int], bottom: tuple[int, int, int]) -> Image.Image:
    img = Image.new("RGB", (width, height), top)
    px = img.load()
    for y in range(height):
        t = y / max(1, height - 1)
        color = tuple(lerp(top[i], bottom[i], t) for i in range(3))
        for x in range(width):
            px[x, y] = color
    return img


def add_soft_grid(draw: ImageDraw.ImageDraw, width: int, height: int) -> None:
    grid_color = (218, 230, 242)
    for x in range(100, width, 120):
        draw.line((x, 0, x, height), fill=grid_color, width=1)
    for y in range(100, height, 120):
        draw.line((0, y, width, y), fill=grid_color, width=1)


def rounded_shadow(base: Image.Image, rect: tuple[int, int, int, int], radius: int, fill: tuple[int, int, int], shadow_alpha: int = 55) -> None:
    shadow = Image.new("RGBA", base.size, (255, 255, 255, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    x0, y0, x1, y1 = rect
    shadow_draw.rounded_rectangle((x0 + 12, y0 + 16, x1 + 12, y1 + 16), radius=radius, fill=(16, 24, 40, shadow_alpha))
    shadow = shadow.filter(ImageFilter.GaussianBlur(18))
    base.alpha_composite(shadow)
    draw = ImageDraw.Draw(base)
    draw.rounded_rectangle(rect, radius=radius, fill=fill, outline=(255, 255, 255, 180), width=3)


def draw_text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill, anchor_y: float = 0.5) -> None:
    x0, y0, x1, y1 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    x = x0 + (x1 - x0 - w) / 2
    y = y0 + (y1 - y0 - h) * anchor_y
    draw.text((x, y), text, font=font, fill=fill)


def draw_multiline(draw: ImageDraw.ImageDraw, x: int, y: int, lines: list[str], font, fill, line_gap: int = 10) -> None:
    cy = y
    for line in lines:
        draw.text((x, cy), line, font=font, fill=fill)
        bbox = draw.textbbox((x, cy), line, font=font)
        cy += (bbox[3] - bbox[1]) + line_gap


def draw_chip(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: tuple[int, int, int], outline: tuple[int, int, int] | None = None) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline or fill, width=2)
    draw_text_center(draw, xy, text, FONT_SMALL, PALETTE["text"])


def draw_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], color: tuple[int, int, int], width: int = 8) -> None:
    if len(points) < 2:
        return
    draw.line(points, fill=color, width=width, joint="curve")
    (x1, y1), (x2, y2) = points[-2], points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_len = 22
    wing = math.pi / 7
    p1 = (x2 - arrow_len * math.cos(angle - wing), y2 - arrow_len * math.sin(angle - wing))
    p2 = (x2 - arrow_len * math.cos(angle + wing), y2 - arrow_len * math.sin(angle + wing))
    draw.polygon([points[-1], p1, p2], fill=color)


def draw_dashed_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: tuple[int, int, int], width: int = 6, dash: int = 16, gap: int = 10) -> None:
    x1, y1 = start
    x2, y2 = end
    distance = math.dist(start, end)
    if distance == 0:
        return
    dx = (x2 - x1) / distance
    dy = (y2 - y1) / distance
    progress = 0
    while progress < distance - dash:
        sx = x1 + dx * progress
        sy = y1 + dy * progress
        ex = x1 + dx * min(progress + dash, distance)
        ey = y1 + dy * min(progress + dash, distance)
        draw.line((sx, sy, ex, ey), fill=color, width=width)
        progress += dash + gap
    draw_arrow(draw, [start, end], color, width=width)


def draw_label_badge(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: tuple[int, int, int], text_fill: tuple[int, int, int] = PALETTE["white"]) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill)
    draw_text_center(draw, xy, text, FONT_SMALL, text_fill)


def draw_swimlane_header(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, title: str, fill: tuple[int, int, int]) -> None:
    draw.rounded_rectangle((x, y, x + w, y + 64), radius=26, fill=fill)
    draw_text_center(draw, (x, y, x + w, y + 64), title, FONT_SMALL, PALETTE["white"])


def draw_architecture_figure(path: Path) -> None:
    width, height = 2200, 1500
    bg = vertical_gradient(width, height, PALETTE["bg_top"], PALETTE["bg_bottom"]).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    # Header
    draw.text((120, 72), "高品质樱桃供应链可信溯源系统", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((122, 150), "四层架构围绕“采集、验签、评分、锚定、查询验证”展开", font=FONT_BODY, fill=PALETTE["muted"])

    layer_specs = [
        ("感知层", "#D9F4EF", ["STM32H743", "SHT31", "MH-Z19B", "ADXL345", "ATECC608A", "ESP8266 / SX1278"]),
        ("接入层", "#E0F2FE", ["HTTPS 上报", "设备认证", "签名信封", "批次标识", "边缘缓冲"]),
        ("服务层", "#E8EEFF", ["签名验证", "规范化哈希", "双层幂等", "品质评分", "告警中心", "Trace / Admin API"]),
        ("存证层", "#EEF2FF", ["PostgreSQL", "Anchor Worker", "Anchor Receipts", "EVM 合约", "链上验证接口"]),
    ]
    x0, w, h, gap = 120, 1640, 205, 38
    start_y = 250
    card_positions: list[tuple[int, int, int, int]] = []
    title_colors = [PALETTE["teal"], PALETTE["cyan"], PALETTE["blue"], (99, 102, 241)]

    for idx, (title, color_hex, chips) in enumerate(layer_specs):
        y = start_y + idx * (h + gap)
        rect = (x0, y, x0 + w, y + h)
        card_positions.append(rect)
        rounded_shadow(bg, rect, radius=36, fill=hex_color(color_hex))
        draw.rounded_rectangle((x0 + 26, y + 24, x0 + 250, y + h - 24), radius=28, fill=title_colors[idx])
        draw_text_center(draw, (x0 + 26, y + 20, x0 + 250, y + h - 20), title, FONT_H2, PALETTE["white"])

        chip_x = x0 + 295
        chip_y = y + 42
        chip_gap_x = 28
        chip_gap_y = 24
        chip_w = 240
        chip_h = 54
        for chip_idx, chip in enumerate(chips):
            row = chip_idx // 3
            col = chip_idx % 3
            cx = chip_x + col * (chip_w + chip_gap_x)
            cy = chip_y + row * (chip_h + chip_gap_y)
            draw_chip(draw, (cx, cy, cx + chip_w, cy + chip_h), chip, PALETTE["white"], outline=PALETTE["slate"])

    # Vertical flow arrows
    for idx in range(3):
        upper = card_positions[idx]
        lower = card_positions[idx + 1]
        start = (upper[0] + upper[2]) // 2, upper[3] + 10
        end = (lower[0] + lower[2]) // 2, lower[1] - 10
        draw_arrow(draw, [start, end], PALETTE["line"], width=10)

    # Right-side user / query lane
    side_rect = (1820, 308, 2070, 1160)
    rounded_shadow(bg, side_rect, radius=36, fill=(255, 255, 255))
    draw.text((1860, 350), "访问角色", font=FONT_H2, fill=PALETTE["text"])
    role_chips = [
        ("管理员", "#DBEAFE"),
        ("操作员", "#CCFBF1"),
        ("监管方", "#FEF3C7"),
        ("消费者", "#FEE2E2"),
    ]
    cy = 460
    for role, fill_hex in role_chips:
        draw_chip(draw, (1860, cy, 2030, cy + 58), role, hex_color(fill_hex), outline=PALETTE["slate"])
        cy += 92

    draw_multiline(
        draw,
        1860,
        850,
        [
            "上行链路",
            "采集  -> 签名  -> 上传",
            "",
            "下行链路",
            "查询  -> 展示  -> 验证",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=14,
    )

    draw_arrow(draw, [(1760, 780), (1820, 780)], PALETTE["blue"], width=8)
    draw_arrow(draw, [(1820, 910), (1760, 910)], PALETTE["teal"], width=8)

    # Footer legend
    draw.rounded_rectangle((120, 1340, 940, 1430), radius=24, fill=(255, 255, 255, 210), outline=PALETTE["slate"], width=2)
    draw.text((150, 1366), "图意说明：数据自上而下完成采集、接入、服务处理与链上存证；用户通过服务层完成查询与验证。", font=FONT_SMALL, fill=PALETTE["muted"])

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def draw_state_machine_figure(path: Path) -> None:
    width, height = 2200, 1450
    bg = vertical_gradient(width, height, PALETTE["bg_top"], (248, 250, 252)).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    draw.text((110, 72), "IngestStatus 事件生命周期", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((112, 150), "以状态机驱动重试、告警与人工回收，保证区块链锚定流程具备最终一致性", font=FONT_BODY, fill=PALETTE["muted"])

    boxes = {
        "received": Box(150, 320, 380, 180, hex_color("#DBEAFE"), "RECEIVED", "已接收 / 待处理"),
        "anchoring": Box(860, 320, 430, 180, hex_color("#E0F2FE"), "ANCHORING", "提交锚定 / 等待确认"),
        "anchored": Box(1540, 250, 420, 180, hex_color("#DCFCE7"), "ANCHORED", "成功终态"),
        "retry": Box(860, 850, 430, 180, hex_color("#FEF3C7"), "FAILED_RETRYING", "可恢复失败 / 等待重试"),
        "dead": Box(150, 980, 420, 180, hex_color("#FEE2E2"), "DEAD_LETTER", "不可恢复失败"),
    }

    for key, box in boxes.items():
        rect = (box.x, box.y, box.x + box.w, box.y + box.h)
        rounded_shadow(bg, rect, radius=36, fill=box.fill)
        accent = {
            "received": PALETTE["blue"],
            "anchoring": PALETTE["cyan"],
            "anchored": PALETTE["green"],
            "retry": PALETTE["amber"],
            "dead": PALETTE["red"],
        }[key]
        draw.rounded_rectangle((box.x + 26, box.y + 24, box.x + 210, box.y + 72), radius=22, fill=accent)
        draw.text((box.x + 56, box.y + 32), box.title, font=FONT_CODE, fill=PALETTE["white"])
        draw.text((box.x + 44, box.y + 100), box.subtitle, font=FONT_BODY, fill=PALETTE["text"])

    # Main transitions
    draw_arrow(draw, [(530, 410), (860, 410)], PALETTE["blue"], width=10)
    draw.text((650, 360), "开始锚定", font=FONT_SMALL, fill=PALETTE["muted"])

    draw_arrow(draw, [(1290, 370), (1540, 340)], PALETTE["green"], width=10)
    draw.text((1370, 305), "成功", font=FONT_SMALL, fill=PALETTE["green"])

    draw_arrow(draw, [(1080, 500), (1080, 850)], PALETTE["amber"], width=10)
    draw.text((1125, 650), "失败且 retry < 3", font=FONT_SMALL, fill=PALETTE["orange"])

    draw_arrow(draw, [(860, 940), (570, 1070)], PALETTE["red"], width=10)
    draw.text((635, 955), "失败且 retry ≥ 3", font=FONT_SMALL, fill=PALETTE["red"])

    draw_dashed_arrow(draw, (360, 980), (360, 500), PALETTE["blue"], width=7)
    draw.text((110, 720), "管理员 requeue", font=FONT_SMALL, fill=PALETTE["blue"])

    # Idempotency note
    note_rect = (1460, 650, 1960, 880)
    rounded_shadow(bg, note_rect, radius=32, fill=(255, 255, 255))
    draw.text((1500, 700), "幂等请求旁路", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        1500,
        770,
        [
            "若请求命中 idempotency_key 或",
            "canonical_hash，系统直接返回",
            "已有记录，不重复进入状态机。",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )
    draw_dashed_arrow(draw, (540, 420), (1460, 760), PALETTE["cyan"], width=6)

    # Legend
    legend_rect = (1210, 1035, 1960, 1220)
    rounded_shadow(bg, legend_rect, radius=28, fill=(255, 255, 255))
    draw.text((1250, 1075), "状态语义", font=FONT_H2, fill=PALETTE["text"])
    legend = [
        ("处理中", PALETTE["blue"]),
        ("成功终态", PALETTE["green"]),
        ("可重试", PALETTE["amber"]),
        ("失败终态", PALETTE["red"]),
    ]
    lx = 1250
    for label, color in legend:
        draw.rounded_rectangle((lx, 1145, lx + 36, 1181), radius=10, fill=color)
        draw.text((lx + 50, 1142), label, font=FONT_SMALL, fill=PALETTE["muted"])
        lx += 170

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def draw_data_flow_figure(path: Path) -> None:
    width, height = 2200, 1480
    bg = vertical_gradient(width, height, PALETTE["bg_top"], (244, 248, 252)).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    draw.text((110, 70), "系统端到端数据流", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((112, 148), "覆盖设备采集、可信签名、后端处理、链上锚定与消费者查询验证两条主链路", font=FONT_BODY, fill=PALETTE["muted"])

    upper_boxes = [
        Box(120, 305, 240, 150, hex_color("#D9F4EF"), "采集节点", "SensorTask\nSHT31 / MH-Z19B / ADXL345"),
        Box(400, 305, 240, 150, hex_color("#E0F2FE"), "签名封装", "SHA-256\nATECC608A\nsignature_envelope"),
        Box(680, 305, 240, 150, hex_color("#E0F2FE"), "无线发送", "ESP8266 / SX1278\nHTTPS / LoRa"),
        Box(960, 305, 280, 150, hex_color("#E8EEFF"), "接入与验签", "FastAPI ingest\n验签 + canonical_hash"),
        Box(1280, 305, 290, 150, hex_color("#EEF2FF"), "业务处理", "幂等检查\n入库 + 品质评分\n告警"),
        Box(1610, 305, 250, 150, hex_color("#E0F2FE"), "锚定引擎", "anchor_worker\nretry_worker"),
        Box(1900, 305, 180, 150, hex_color("#DCFCE7"), "链上存证", "EVM 合约\nverifyAnchor"),
    ]

    accents = [PALETTE["teal"], PALETTE["cyan"], PALETTE["cyan"], PALETTE["blue"], (99, 102, 241), PALETTE["orange"], PALETTE["green"]]
    for idx, box in enumerate(upper_boxes):
        rect = (box.x, box.y, box.x + box.w, box.y + box.h)
        rounded_shadow(bg, rect, radius=34, fill=box.fill)
        draw_label_badge(draw, (box.x + 20, box.y + 20, box.x + 128, box.y + 60), f"S{idx+1}", accents[idx])
        draw.text((box.x + 24, box.y + 78), box.title, font=FONT_H2, fill=PALETTE["text"])
        lines = box.subtitle.split("\n")
        draw_multiline(draw, box.x + 26, box.y + 125, lines, FONT_SMALL, PALETTE["muted"], line_gap=8)

    for idx in range(len(upper_boxes) - 1):
        left = upper_boxes[idx]
        right = upper_boxes[idx + 1]
        start = (left.x + left.w, left.y + left.h // 2)
        end = (right.x, right.y + right.h // 2)
        draw_arrow(draw, [start, end], PALETTE["line"], width=8)

    draw.rounded_rectangle((120, 560, 2080, 590), radius=15, fill=(255, 255, 255, 170), outline=PALETTE["slate"], width=1)
    draw.text((140, 562), "上行数据流：传感器采集 -> 可信签名 -> 无线传输 -> 验签入库 -> 评分告警 -> 区块链锚定", font=FONT_SMALL, fill=PALETTE["muted"])

    lower_boxes = [
        Box(210, 820, 240, 150, hex_color("#FEE2E2"), "消费者", "扫码 / 输入 batch_id"),
        Box(520, 820, 280, 150, hex_color("#E0F2FE"), "公开查询接口", "GET /v1/public/trace/{batch_id}"),
        Box(870, 820, 260, 150, hex_color("#E8EEFF"), "查询聚合", "事件时间线\n品质等级\ntx_hash"),
        Box(1200, 820, 280, 150, hex_color("#EEF2FF"), "前端展示", "时间线 / 曲线图\n可视化详情"),
        Box(1550, 820, 290, 150, hex_color("#DCFCE7"), "独立验证", "重新计算 canonical_hash\n调用链上 verifyAnchor"),
    ]
    lower_accents = [PALETTE["red"], PALETTE["cyan"], PALETTE["blue"], PALETTE["teal"], PALETTE["green"]]
    for idx, box in enumerate(lower_boxes):
        rect = (box.x, box.y, box.x + box.w, box.y + box.h)
        rounded_shadow(bg, rect, radius=32, fill=box.fill)
        draw_label_badge(draw, (box.x + 18, box.y + 18, box.x + 126, box.y + 58), f"Q{idx+1}", lower_accents[idx])
        draw.text((box.x + 22, box.y + 76), box.title, font=FONT_H2, fill=PALETTE["text"])
        draw_multiline(draw, box.x + 24, box.y + 124, box.subtitle.split("\n"), FONT_SMALL, PALETTE["muted"], line_gap=8)

    for idx in range(len(lower_boxes) - 1):
        left = lower_boxes[idx]
        right = lower_boxes[idx + 1]
        start = (left.x + left.w, left.y + left.h // 2)
        end = (right.x, right.y + right.h // 2)
        draw_arrow(draw, [start, end], PALETTE["line"], width=8)

    draw_dashed_arrow(draw, (1470, 430), (910, 810), PALETTE["cyan"], width=6)
    draw.text((1080, 610), "查询时读取数据库与链上凭证", font=FONT_SMALL, fill=PALETTE["cyan"])

    draw.rounded_rectangle((1540, 1065, 2080, 1300), radius=28, fill=(255, 255, 255, 215), outline=PALETTE["slate"], width=2)
    draw.text((1570, 1105), "数据完整性闭环", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        1570,
        1170,
        [
            "1. 设备侧计算摘要并硬件签名",
            "2. 后端生成 canonical_hash 并入库",
            "3. 仅将哈希锚定至链上",
            "4. 查询阶段可由用户独立复核",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def draw_hash_flow_figure(path: Path) -> None:
    width, height = 2200, 1480
    bg = vertical_gradient(width, height, (248, 250, 252), PALETTE["bg_bottom"]).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    draw.text((110, 72), "哈希规范化流程", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((112, 148), "通过统一的序列化规则，确保 Python 与 TypeScript 对同一事件得到完全一致的 canonical_hash", font=FONT_BODY, fill=PALETTE["muted"])

    input_rect = (120, 300, 560, 1110)
    rounded_shadow(bg, input_rect, radius=34, fill=(255, 255, 255))
    draw_label_badge(draw, (150, 330, 300, 380), "输入事件 JSON", PALETTE["blue"])
    draw_multiline(
        draw,
        155,
        430,
        [
            "{",
            '  "device_id": "node-01",',
            '  "batch_id": "CH2026-0008",',
            '  "temperature": 2.1,',
            '  "humidity": 88,',
            '  "co2": 620,',
            '  "timestamp": "2026-03-28T08:30:00+08:00",',
            '  "meta": { "location": "冷库A", "stage": "warehouse" }',
            "}",
        ],
        FONT_SMALL,
        PALETTE["text"],
        line_gap=12,
    )
    draw.text((155, 930), "问题：不同语言默认序列化、时区与空白处理不一致，会导致哈希值不稳定。", font=FONT_SMALL, fill=PALETTE["muted"])

    steps = [
        ("01", "字段递归排序", "对象键按字典序稳定排序"),
        ("02", "时间标准化", "统一转为 UTC ISO8601"),
        ("03", "字符串 trim", "去除前后空白与无效换行"),
        ("04", "ASCII 转义", "非 ASCII 字符统一 \\uXXXX"),
        ("05", "紧凑序列化", "使用无空格分隔符"),
        ("06", "SHA-256 哈希", "得到 64 位十六进制摘要"),
    ]
    step_w, step_h = 280, 128
    sx, sy = 700, 310
    for idx, (num, title, sub) in enumerate(steps):
        row = idx // 2
        col = idx % 2
        x = sx + col * 340
        y = sy + row * 210
        fill = [hex_color("#E0F2FE"), hex_color("#EEF2FF"), hex_color("#D9F4EF"), hex_color("#FEF3C7"), hex_color("#E8EEFF"), hex_color("#DCFCE7")][idx]
        accent = [PALETTE["blue"], (99, 102, 241), PALETTE["teal"], PALETTE["amber"], PALETTE["cyan"], PALETTE["green"]][idx]
        rounded_shadow(bg, (x, y, x + step_w, y + step_h), radius=28, fill=fill)
        draw_label_badge(draw, (x + 18, y + 18, x + 92, y + 56), num, accent)
        draw.text((x + 116, y + 20), title, font=FONT_H2, fill=PALETTE["text"])
        draw.text((x + 24, y + 78), sub, font=FONT_SMALL, fill=PALETTE["muted"])

    arrow_points = [
        ((980, 374), (1040, 374)),
        ((1380, 374), (1380, 520), (1040, 520)),
        ((980, 584), (1040, 584)),
        ((1380, 584), (1380, 730), (1040, 730)),
        ((980, 794), (1040, 794)),
    ]
    for pts in arrow_points:
        draw_arrow(draw, list(pts), PALETTE["line"], width=7)

    output_rect = (1620, 300, 2080, 1110)
    rounded_shadow(bg, output_rect, radius=34, fill=(255, 255, 255))
    draw_label_badge(draw, (1650, 330, 1840, 380), "输出结果", PALETTE["green"])
    draw.text((1655, 440), "canonical_hash", font=FONT_H2, fill=PALETTE["text"])
    hash_lines = [
        "9f4f3bb18f1b0a9d...",
        "3e1204b7f5a6c261...",
        "1f8f4fd22fb6d7aa",
    ]
    draw_multiline(draw, 1658, 520, hash_lines, FONT_CODE, PALETTE["green"], line_gap=10)
    draw.text((1655, 760), "一致性目标", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        1658,
        830,
        [
            "Python 后端 = TypeScript 前端",
            "相同事件 -> 相同哈希",
            "哈希可直接用于幂等与链上锚定",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )

    draw.rounded_rectangle((710, 1230, 1500, 1350), radius=24, fill=(255, 255, 255, 220), outline=PALETTE["slate"], width=2)
    draw.text((740, 1268), "规范化的核心价值：消除跨语言、跨平台、跨时区带来的哈希漂移，使链上存证与前端自验具备可重复性。", font=FONT_SMALL, fill=PALETTE["muted"])

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def draw_idempotency_sequence_figure(path: Path) -> None:
    width, height = 2200, 1480
    bg = vertical_gradient(width, height, PALETTE["bg_top"], (245, 249, 252)).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    draw.text((110, 72), "双层幂等处理时序", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((112, 148), "第一层防重放，第二层防内容重复；依赖数据库约束与回查策略处理并发竞争", font=FONT_BODY, fill=PALETTE["muted"])

    lanes = [
        ("客户端", 160, 280, PALETTE["teal"]),
        ("API / ingest", 560, 320, PALETTE["blue"]),
        ("ingest_requests", 980, 340, PALETTE["cyan"]),
        ("events", 1410, 250, (99, 102, 241)),
        ("响应", 1780, 250, PALETTE["green"]),
    ]
    line_top, line_bottom = 280, 1250
    lane_centers = []
    for title, x, w, fill in lanes:
        draw_swimlane_header(draw, x, 225, w, title, fill)
        cx = x + w // 2
        lane_centers.append(cx)
        draw.line((cx, line_top, cx, line_bottom), fill=PALETTE["line"], width=4)

    steps = [
        (330, 0, 1, "POST /v1/events + idempotency_key", PALETTE["blue"]),
        (420, 1, 2, "查询 idempotency_key 是否存在", PALETTE["cyan"]),
        (510, 2, 1, "不存在 -> 继续处理", PALETTE["cyan"]),
        (600, 1, 3, "计算 canonical_hash", PALETTE["blue"]),
        (690, 1, 3, "尝试 INSERT events(canonical_hash)", (99, 102, 241)),
        (780, 3, 1, "写入成功", PALETTE["green"]),
        (870, 1, 2, "记录 idempotency_key -> event_id", PALETTE["cyan"]),
        (960, 1, 4, "返回 201 Created", PALETTE["green"]),
    ]
    for y, from_idx, to_idx, label, color in steps:
        draw_arrow(draw, [(lane_centers[from_idx], y), (lane_centers[to_idx], y)], color, width=6)
        draw.text((((lane_centers[from_idx] + lane_centers[to_idx]) // 2) - 140, y - 34), label, font=FONT_SMALL, fill=PALETTE["muted"])

    # duplicate branch
    branch_rect = (118, 1010, 960, 1290)
    rounded_shadow(bg, branch_rect, radius=30, fill=(255, 255, 255))
    draw.text((150, 1045), "重复请求分支", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        150,
        1110,
        [
            "A. 若命中 idempotency_key：直接返回已有 event_id",
            "B. 若未命中但 canonical_hash 冲突：回滚后查询已有记录",
            "C. 因此可同时防止“重放请求”和“内容重复上传”",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )
    draw_dashed_arrow(draw, (lane_centers[2], 420), (960, 1090), PALETTE["amber"], width=6)
    draw_dashed_arrow(draw, (lane_centers[3], 690), (960, 1160), PALETTE["orange"], width=6)

    conflict_rect = (1090, 1010, 1960, 1290)
    rounded_shadow(bg, conflict_rect, radius=30, fill=(255, 255, 255))
    draw.text((1125, 1045), "并发安全要点", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        1125,
        1110,
        [
            "依赖 UNIQUE(idempotency_key) 与 UNIQUE(canonical_hash)",
            "采用“乐观写入 + IntegrityError 捕获 + 回查”",
            "避免先查后写造成的竞态窗口",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def draw_anchor_sequence_figure(path: Path) -> None:
    width, height = 2200, 1480
    bg = vertical_gradient(width, height, (247, 250, 252), PALETTE["bg_bottom"]).convert("RGBA")
    draw = ImageDraw.Draw(bg)
    add_soft_grid(draw, width, height)

    draw.text((110, 72), "区块链锚定引擎工作时序", font=FONT_TITLE, fill=PALETTE["text"])
    draw.text((112, 148), "通过 anchor_worker、retry_worker 和提交记录表协同，实现崩溃恢复、确认轮询与失败重试", font=FONT_BODY, fill=PALETTE["muted"])

    lanes = [
        ("anchor_worker", 130, 280, PALETTE["blue"]),
        ("数据库", 510, 250, (99, 102, 241)),
        ("AnchorAdapter", 870, 300, PALETTE["cyan"]),
        ("EVM 节点", 1280, 240, PALETTE["teal"]),
        ("retry_worker", 1650, 280, PALETTE["orange"]),
    ]
    centers = []
    for title, x, w, fill in lanes:
        draw_swimlane_header(draw, x, 225, w, title, fill)
        cx = x + w // 2
        centers.append(cx)
        draw.line((cx, 290, cx, 1240), fill=PALETTE["line"], width=4)

    seq_steps = [
        (340, 0, 1, "查询 RECEIVED 事件", PALETTE["blue"]),
        (430, 1, 0, "返回待锚定事件", (99, 102, 241)),
        (520, 0, 1, "写入 PENDING submission", PALETTE["blue"]),
        (610, 0, 2, "adapter.anchor_event()", PALETTE["cyan"]),
        (700, 2, 3, "提交交易 / 获取 tx_hash", PALETTE["teal"]),
        (790, 3, 2, "返回 tx_hash", PALETTE["teal"]),
        (880, 2, 1, "保存 tx_hash + status=PENDING", PALETTE["cyan"]),
        (970, 0, 2, "轮询 get_receipt()", PALETTE["blue"]),
        (1060, 2, 3, "查询 receipt / confirmations", PALETTE["teal"]),
        (1150, 3, 2, "返回 receipt", PALETTE["teal"]),
        (1240, 2, 1, "FINALIZED 或 REORGED", PALETTE["green"]),
    ]
    for y, from_idx, to_idx, label, color in seq_steps:
        draw_arrow(draw, [(centers[from_idx], y), (centers[to_idx], y)], color, width=6)
        draw.text((((centers[from_idx] + centers[to_idx]) // 2) - 150, y - 34), label, font=FONT_SMALL, fill=PALETTE["muted"])

    retry_rect = (1540, 610, 2040, 1020)
    rounded_shadow(bg, retry_rect, radius=32, fill=(255, 255, 255))
    draw.text((1575, 655), "失败重试分支", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        1575,
        730,
        [
            "1. 锚定异常 -> 事件状态置为 FAILED_RETRYING",
            "2. retry_worker 从重试队列取出任务",
            "3. retry_count < 3 时再次提交",
            "4. 超过阈值则转入 DEAD_LETTER 并告警",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )
    draw_dashed_arrow(draw, (centers[1], 1240), (centers[4], 760), PALETTE["amber"], width=6)

    recovery_rect = (140, 1080, 1160, 1290)
    rounded_shadow(bg, recovery_rect, radius=30, fill=(255, 255, 255))
    draw.text((175, 1122), "崩溃恢复机制", font=FONT_H2, fill=PALETTE["text"])
    draw_multiline(
        draw,
        175,
        1188,
        [
            "系统在提交交易前先写入 PENDING submission record。",
            "若服务崩溃重启，可继续轮询同一 transaction_hash，",
            "避免重复上链并保持锚定流程幂等。",
        ],
        FONT_SMALL,
        PALETTE["muted"],
        line_gap=12,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    bg.convert("RGB").save(path, quality=95)


def insert_paragraph_before(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None  # type: ignore[assignment]


def set_center(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = 0
    fmt.space_after = 0


def replace_placeholder_image(doc_path: Path, output_path: Path, replacements: dict[str, tuple[Path, float]]) -> None:
    shutil.copy2(doc_path, output_path)
    doc = Document(output_path)
    paragraphs = list(doc.paragraphs)
    for caption, (image_path, width_cm) in replacements.items():
        target_idx = None
        for idx, para in enumerate(paragraphs):
            if para.text.strip().replace("\n", " ") == caption:
                target_idx = idx
                break
        if target_idx is None:
            continue

        caption_para = doc.paragraphs[target_idx]
        pic_para = doc.paragraphs[target_idx - 1]
        if pic_para.runs:
            remove_paragraph(pic_para)
            caption_para = doc.paragraphs[target_idx - 1]

        new_pic_para = insert_paragraph_before(caption_para)
        set_center(new_pic_para)
        new_pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))

    doc.save(output_path)


def main() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    arch_path = FIGURE_DIR / "figure_3_1_architecture.png"
    data_flow_path = FIGURE_DIR / "figure_3_2_data_flow.png"
    state_path = FIGURE_DIR / "figure_3_4_state_machine.png"
    hash_flow_path = FIGURE_DIR / "figure_5_2_hash_flow.png"
    idempotency_path = FIGURE_DIR / "figure_5_3_idempotency_sequence.png"
    anchor_path = FIGURE_DIR / "figure_5_5_anchor_sequence.png"

    draw_architecture_figure(arch_path)
    draw_data_flow_figure(data_flow_path)
    draw_state_machine_figure(state_path)
    draw_hash_flow_figure(hash_flow_path)
    draw_idempotency_sequence_figure(idempotency_path)
    draw_anchor_sequence_figure(anchor_path)

    replace_placeholder_image(
        INPUT_DOC,
        DRAFT_DOC,
        {
            "图3-1 系统分层架构图": (arch_path, 15.8),
            "图3-2 系统端到端数据流图": (data_flow_path, 16.0),
            "图3-4 IngestStatus状态机图": (state_path, 15.2),
            "图5-2 哈希规范化流程图": (hash_flow_path, 15.6),
            "图5-3 双层幂等处理时序图": (idempotency_path, 16.0),
            "图5-5 区块链锚定引擎工作时序图": (anchor_path, 16.0),
        },
    )

    print(arch_path)
    print(data_flow_path)
    print(state_path)
    print(hash_flow_path)
    print(idempotency_path)
    print(anchor_path)
    print(DRAFT_DOC)


if __name__ == "__main__":
    main()
