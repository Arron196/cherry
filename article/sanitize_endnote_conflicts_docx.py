from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
SOURCE_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v18_扩展引文_字段刷新.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v19_EndNote安全版.docx"


PARAGRAPH_REPLACEMENTS = {
    r"计算消息摘要 e = \text{SHA-256}(m)（本系统中由STM32在发送给ATECC608A前完成）":
        "计算消息摘要 e = SHA-256(m)（本系统中由STM32在发送给ATECC608A前完成）",
    "消费者扫描溯源二维码，访问GET /v1/public/trace/{batch_id}":
        "消费者扫描溯源二维码，访问 GET /v1/public/trace/<batch_id>",
    "GET /v1/public/trace/{batch_id}（公开溯源查询接口）：":
        "GET /v1/public/trace/<batch_id>（公开溯源查询接口）：",
    "规则1：字典键递归升序排序 所有JSON对象（包括嵌套对象）的键按Unicode码点升序排列。例如：{\"b\": 1, \"a\": 2} 规范化为 {\"a\": 2, \"b\": 1}。":
        "规则1：字典键递归升序排序 所有JSON对象（包括嵌套对象）的键按Unicode码点升序排列。例如，对象“b: 1, a: 2”规范化后写为“a: 2, b: 1”。",
    "系统维护一套共享测试向量（Shared Test Vectors），包含若干组{输入JSON, 预期哈希值}对，在Python单元测试（pytest）和前端测试（Vitest）中分别验证两个实现的输出完全一致。关键测试向量覆盖：非ASCII字符（中文设备名）、嵌套对象、datetime时区转换（UTC+8 → UTC+0）、浮点数精度（Python float vs TypeScript number）等边界情况。":
        "系统维护一套共享测试向量（Shared Test Vectors），包含若干组“输入 JSON 与预期哈希值”的测试对，在Python单元测试（pytest）和前端测试（Vitest）中分别验证两个实现的输出完全一致。关键测试向量覆盖：非ASCII字符（中文设备名）、嵌套对象、datetime时区转换（UTC+8 → UTC+0）、浮点数精度（Python float vs TypeScript number）等边界情况。",
    "设备端固件为每次发送请求生成一个全局唯一的幂等性键（idempotency_key），格式为 {device_id}:{timestamp_ms}:{random_nonce}。该键通过HTTP请求头 Idempotency-Key 随请求发送至后端。":
        "设备端固件为每次发送请求生成一个全局唯一的幂等性键（idempotency_key），格式为 <device_id>:<timestamp_ms>:<random_nonce>。该键通过HTTP请求头 Idempotency-Key 随请求发送至后端。",
    'rollout_canary_outcomes_total{outcome="success|failure"} — Canary成功/失败计数':
        "rollout_canary_outcomes_total（其中 outcome 标签取 success 或 failure）— Canary成功/失败计数",
}


CELL_REPLACEMENTS = {
    "GET /v1/public/trace/{id}": "GET /v1/public/trace/<id>",
    "GET /v1/trace/{id}": "GET /v1/trace/<id>",
}


def replace_text_in_paragraph(paragraph) -> int:
    text = paragraph.text
    new_text = text
    for old, new in PARAGRAPH_REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    new_text = new_text.replace("frame = {}", "frame = dict()")
    if new_text == text:
        return 0

    if not paragraph.runs:
        paragraph.text = new_text
        return 1

    first = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first.text = new_text
    return 1


def replace_text_in_cell(cell) -> int:
    changed = 0
    for paragraph in cell.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in CELL_REPLACEMENTS.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            if paragraph.runs:
                first = paragraph.runs[0]
                for run in paragraph.runs[1:]:
                    run.text = ""
                first.text = new_text
            else:
                paragraph.text = new_text
            changed += 1
    return changed


def main() -> None:
    shutil.copy2(SOURCE_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    changed = 0
    for paragraph in doc.paragraphs:
        changed += replace_text_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += replace_text_in_cell(cell)

    doc.save(OUTPUT_DOC)
    print(f"changed={changed}")
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
