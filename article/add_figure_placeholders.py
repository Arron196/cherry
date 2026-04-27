from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_插图占位版_草稿.docx"
ASSET_DIR = ROOT / "figure_placeholders"


@dataclass
class PlaceholderSpec:
    anchor_text: str
    caption: str
    note_lines: list[str]
    width_cm: float = 12.5
    height_px: int = 920
    remove_next_ascii: bool = False


PLACEHOLDERS: list[PlaceholderSpec] = [
    PlaceholderSpec(
        anchor_text="系统采用四层架构：感知层、接入层、服务层、存证层，各层职责清晰，通过标准接口解耦，使各层可独立演进。",
        caption="图3-1 系统分层架构图",
        note_lines=[
            "建议插入：四层系统架构图",
            "内容应包含感知层、接入层、服务层、存证层",
            "可用 draw.io / Visio / ProcessOn 绘制后替换本图",
        ],
        remove_next_ascii=True,
    ),
    PlaceholderSpec(
        anchor_text="3.2.2 数据流设计",
        caption="图3-2 系统端到端数据流图",
        note_lines=[
            "建议插入：采集 -> 签名 -> 上传 -> 入库 -> 评分 -> 锚定 -> 查询验证",
            "推荐画成时序式或泳道式流程图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="3.3.1 核心数据表结构",
        caption="图3-3 核心数据表ER关系图",
        note_lines=[
            "建议插入：events / ingest_requests / anchor_submission_records / quality_results 等核心表关系图",
            "可突出主键、外键和业务流向",
        ],
    ),
    PlaceholderSpec(
        anchor_text="事件的生命周期由IngestStatus状态机管理，共5个状态：",
        caption="图3-4 IngestStatus状态机图",
        note_lines=[
            "Mermaid 图应导出为 PNG 或 SVG 后插入",
            "状态建议包含 RECEIVED / ANCHORING / ANCHORED / FAILED_RETRYING / DEAD_LETTER",
        ],
        remove_next_ascii=True,
    ),
    PlaceholderSpec(
        anchor_text="3.4 系统安全架构",
        caption="图3-5 系统安全架构图",
        note_lines=[
            "建议插入：JWT、RBAC、签名验证、设备密钥管理、审计日志之间的关系图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="4.1 STM32H743主控电路设计",
        caption="图4-1 STM32H743最小系统原理图",
        note_lines=[
            "建议插入：主控最小系统原理图或核心PCB模块图",
            "重点标出晶振、复位、电源去耦、SWD接口",
        ],
    ),
    PlaceholderSpec(
        anchor_text="4.2 传感器接口电路",
        caption="图4-2 多传感器接口与总线连接图",
        note_lines=[
            "建议插入：SHT31、MH-Z19B、ADXL345 与主控的连接示意图",
            "可同时标出 I2C / UART / GPIO 连接关系",
        ],
    ),
    PlaceholderSpec(
        anchor_text="4.3 ATECC608A安全芯片接口与签名流程",
        caption="图4-3 ATECC608A连接与签名流程图",
        note_lines=[
            "建议插入：ATECC608A 与 STM32 的硬件连接图",
            "可叠加签名数据流说明",
        ],
    ),
    PlaceholderSpec(
        anchor_text="4.4 无线通信模块设计",
        caption="图4-4 Wi-Fi/LoRa通信组网图",
        note_lines=[
            "建议插入：ESP8266 / SX1278 与后端服务的组网示意图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="4.5 电源管理设计",
        caption="图4-5 电源管理原理图",
        note_lines=[
            "建议插入：5V 输入、3.3V 稳压、模拟/数字电源隔离示意图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.1 嵌入式固件设计（FreeRTOS）",
        caption="图5-1 FreeRTOS任务调度与数据通路图",
        note_lines=[
            "建议插入：SensorTask / SignTask / CommTask 的任务关系图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.2 哈希规范化算法实现",
        caption="图5-2 哈希规范化流程图",
        note_lines=[
            "建议插入：排序、时间标准化、trim、转义、序列化、哈希计算流程图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.3 双层幂等性防护机制",
        caption="图5-3 双层幂等处理时序图",
        note_lines=[
            "建议插入：客户端、后端、数据库之间的幂等处理时序图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.4.1 验证流程设计",
        caption="图5-4 签名验证流程图",
        note_lines=[
            "建议插入：signature_envelope 解析、密钥查找、哈希重算、验签、结果处理流程",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.6.1 锚定引擎架构",
        caption="图5-5 区块链锚定引擎工作时序图",
        note_lines=[
            "建议插入：anchor_worker / retry_worker / adapter / EVM 节点时序图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.6.2 四阶段渐进上线策略（Rollout）",
        caption="图5-6 四阶段Rollout策略图",
        note_lines=[
            "建议插入：rollback_safe -> shadow -> canary -> full 的阶段切换图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="5.7 前端可视化实现",
        caption="图5-7 前端页面结构与关键界面示意图",
        note_lines=[
            "建议插入：管理员看板、公开溯源查询页、时间线页面截图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="6.1.1 测试环境与方法",
        caption="图6-1 硬件测试平台实物图",
        note_lines=[
            "建议插入：硬件节点、传感器、供电、调试器组成的测试平台照片",
        ],
    ),
    PlaceholderSpec(
        anchor_text="6.4.1 端到端溯源链路测试",
        caption="图6-2 端到端溯源链路验证截图",
        note_lines=[
            "建议插入：从设备上传到前端查询结果的联调截图",
        ],
    ),
    PlaceholderSpec(
        anchor_text="6.4.2 Canary上线SLO监控测试",
        caption="图6-3 Canary监控与自动回滚结果图",
        note_lines=[
            "建议插入：SLO 面板、告警触发和自动回滚结果截图",
        ],
    ),
]


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        ("C:/Windows/Fonts/simhei.ttf", bold),
        ("C:/Windows/Fonts/simsun.ttc", bold),
        ("C:/Windows/Fonts/msyh.ttc", bold),
    ]
    for font_path, _ in candidates:
        if Path(font_path).exists():
            return ImageFont.truetype(font_path, size=size)
    return ImageFont.load_default()


def create_placeholder_image(path: Path, caption: str, note_lines: list[str], height_px: int) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    width_px = 1600
    img = Image.new("RGB", (width_px, height_px), "white")
    draw = ImageDraw.Draw(img)
    border = 12
    draw.rectangle((border, border, width_px - border, height_px - border), outline="black", width=4)

    title_font = get_font(54, bold=True)
    body_font = get_font(36)
    small_font = get_font(30)

    lines = [caption, "", "图片占位"] + note_lines + ["", "后续请替换为真实导出的 PNG / SVG / 实物截图"]
    fonts = [title_font, title_font, body_font, body_font, body_font, body_font, small_font, small_font]

    total_height = 0
    measured: list[tuple[str, tuple[int, int], ImageFont.ImageFont | ImageFont.FreeTypeFont]] = []
    for idx, line in enumerate(lines):
        font = fonts[min(idx, len(fonts) - 1)]
        bbox = draw.textbbox((0, 0), line, font=font)
        size = (bbox[2] - bbox[0], bbox[3] - bbox[1])
        measured.append((line, size, font))
        total_height += size[1] + (16 if line else 10)

    y = max(80, (height_px - total_height) // 2)
    for line, size, font in measured:
        x = (width_px - size[0]) // 2
        draw.text((x, y), line, fill="black", font=font)
        y += size[1] + (16 if line else 10)

    img.save(path)


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None  # type: ignore[assignment]


def format_caption(paragraph: Paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.bold = False
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Times New Roman")
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "Times New Roman")
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")


def format_spacing(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def insert_placeholder_after(anchor: Paragraph, image_path: Path, caption: str, width_cm: float) -> Paragraph:
    pic_para = insert_paragraph_after(anchor)
    format_spacing(pic_para)
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))

    caption_para = insert_paragraph_after(pic_para)
    format_caption(caption_para, caption)

    spacer = insert_paragraph_after(caption_para)
    format_spacing(spacer)
    return spacer


def main() -> None:
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    asset_map: dict[str, Path] = {}
    for idx, spec in enumerate(PLACEHOLDERS, start=1):
        image_path = ASSET_DIR / f"placeholder_{idx:02d}.png"
        create_placeholder_image(image_path, spec.caption, spec.note_lines, spec.height_px)
        asset_map[spec.caption] = image_path

    used_specs: set[int] = set()
    paragraphs = list(doc.paragraphs)
    for para in paragraphs:
        text = para.text.strip().replace("\n", " ")
        for idx, spec in enumerate(PLACEHOLDERS):
            if idx in used_specs:
                continue
            if text == spec.anchor_text:
                last = insert_placeholder_after(para, asset_map[spec.caption], spec.caption, spec.width_cm)
                used_specs.add(idx)
                if spec.remove_next_ascii:
                    next_p = last._p.getnext()
                    while next_p is not None:
                        maybe = Paragraph(next_p, para._parent)
                        next_text = maybe.text.strip().replace("\n", " ")
                        if next_text:
                            if next_text.startswith("┌") or next_text.startswith("RECEIVED "):
                                remove_paragraph(maybe)
                            break
                        next_p = next_p.getnext()
                break

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
