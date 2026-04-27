from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v17.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v18_扩展引文.docx"
SOURCE_MD = ROOT / "论文.md"


TARGET_121 = (
    "国内方面，谭砂文等（2021）[3]综述了区块链在中国农业供应链中的应用进展，"
    "指出目前国内以联盟链为主的溯源方案虽降低了公链的高成本问题，但联盟链的节点管理权仍集中于少数机构，"
    "去中心化程度不足，难以真正消除对运营主体的信任依赖。雷据联盟（2019）[4]在西藏红酒溯源项目中实现了基于"
    "Hyperledger Fabric的供应链记录链上存储，但未集成传感器实时监测功能。"
)

TARGET_122 = (
    "Benkhaddra等（2022）[11]则将区块链溯源技术应用于智能制造场景，证明了该技术路线在高频采集环境（50Hz）下的可行性，"
    "其提出的批量聚合哈希方法有效降低了链上写入频率。Wang等（2023）[12]的冷链溯源系统采用联邦学习对品质预测模型进行分布式训练，"
    "与本文的规则化品质评分体系形成互补。"
)

NEW_PARAS_121 = [
    (
        "近五年的研究进一步推动了农产品溯源从概念验证走向面向具体品类的系统实现。"
        "Mirikar等（2025）[28]、Yang等（2024）[31]、Li等（2024）[37]和Bosona等（2023）[55]分别从IoT感知追溯、"
        "可编辑数据管理、RFID融合和综述分析角度完善了农产品区块链追溯体系；Rahman等（2025）[27]、Baku等（2024）[33]、"
        "Hasan等（2023）[44]和Guan等（2023）[43]则从生鲜品质分析、框架设计、点对点可信架构与模型抽象层面证明，"
        "区块链可显著提升批次信息可验证性与跨主体协同效率。"
    ),
    (
        "在具体业务与质量安全管理场景方面，Indap等（2023）[48]以樱桃产品为案例验证了区块链在高价值果品供应链中的适配性，"
        "Zheng等（2023）[54]、Cai（2023）[46]、Hu（2023）[53]和Meng（2022）[59]则分别从质量追溯、品质安全管理、"
        "系统数字化升级和追溯数据治理角度补充了面向农产品质量安全的系统设计。与此同时，Yu等（2024）[34]、Zhang等（2022）[36]、"
        "Tang等（2023）[49]、Krishna等（2022）[60]、Wisessing等（2022）[62]和Xie等（2022）[63]表明，冷链物流、"
        "冷链食品、电商农产品与物联网食品监测等多场景正在推动追溯系统从“信息记录”向“质量监测+过程管控”演进。"
    ),
]

NEW_PARAS_122 = [
    (
        "从IoT与区块链融合的系统架构看，Subashini等（2022）[58]、Kumar等（2023）[47]、Huang等（2021）[66]、"
        "Guo等（2021）[65]和Nair等（2024）[36]分别从IoT协同框架、智慧农业业务闭环、5G-IoT接入、轻量化农业追溯平台和跨生态收敛架构"
        "等方面拓展了区块链与感知网络的融合路径；Abishek等（2025）[25]、AlRossais等（2025）[30]、Mpyana Mwamba等（2025）[26]、"
        "Chen等（2021）[70]和Marchese等（2021）[64]则进一步强调了运输节点身份管理、链上凭证与平台化部署对系统落地的重要性。"
    ),
    (
        "此外，Bhaskar等（2024）[35]、Wang等（2024）[38]和Feng等（2024）[41]表明，联邦学习、双链分层和机器学习质量预测正成为新一代追溯系统的"
        "重要增强能力；Kumar等（2023）[52]、Pasha等（2023）[50]、Eghmazi等（2024）[39]和Balamurugan等（2021）[67]则从安全供应链、"
        "NB-IoT部署、IoT数据完整性和低资源环境可信接入等角度推进了融合系统的工程化。与此同时，Vitaskos等（2024）[32]、Rajput等（2023）[45]、"
        "Musthafa Sheriff等（2024）[40]、Manoj等（2023）[42]、Jannat等（2021）[68]和Chen等（2021）[69]说明，区块链追溯已经从农产品延伸到"
        "橄榄油、有机食品、多品类食品和利润优化管理等更广泛场景，但在设备侧数据源可信、嵌入式低资源适配和链上成本控制方面仍缺乏端到端闭环实现。"
    ),
]


def parse_reference_lines(markdown_path: Path) -> list[str]:
    text = markdown_path.read_text(encoding="utf-8")
    marker = "## 参考文献"
    idx = text.find(marker)
    if idx < 0:
        raise RuntimeError("参考文献章节未找到")
    block = text[idx + len(marker) :].strip()
    refs = []
    for line in block.splitlines():
        clean = line.strip()
        if clean.startswith("["):
            refs.append(clean.replace("*", ""))
    return refs


def copy_paragraph_style(src: Paragraph, dst: Paragraph) -> None:
    dst.style = src.style
    spf = src.paragraph_format
    dpf = dst.paragraph_format
    dpf.alignment = spf.alignment
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing
    dpf.line_spacing_rule = spf.line_spacing_rule


def copy_run_style(src: Paragraph, dst: Paragraph, text: str) -> None:
    run = dst.add_run(text)
    if src.runs:
        tmpl = src.runs[0]
        run.bold = tmpl.bold
        run.italic = tmpl.italic
        run.underline = tmpl.underline
        run.font.name = tmpl.font.name
        run.font.size = tmpl.font.size
        run.font.color.rgb = tmpl.font.color.rgb
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    copy_paragraph_style(template, new_para)
    copy_run_style(template, new_para, text)
    return new_para


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text == text:
            return para
    raise RuntimeError(f"未找到目标段落: {text[:60]}")


def remove_following_paragraphs(heading_para: Paragraph) -> list[Paragraph]:
    body = heading_para._p.getparent()
    removed = []
    node = heading_para._p.getnext()
    while node is not None:
        next_node = node.getnext()
        para = Paragraph(node, heading_para._parent)
        removed.append(para)
        body.remove(node)
        node = next_node
    return removed


def main() -> None:
    shutil.copy2(SOURCE_DOC, OUTPUT_DOC)
    references = parse_reference_lines(SOURCE_MD)

    doc = Document(OUTPUT_DOC)

    anchor = find_paragraph(doc, TARGET_121)
    for para_text in NEW_PARAS_121:
        anchor = insert_paragraph_after(anchor, para_text, anchor)

    anchor = find_paragraph(doc, TARGET_122)
    template = anchor
    for para_text in NEW_PARAS_122:
        anchor = insert_paragraph_after(anchor, para_text, template)

    ref_heading = find_paragraph(doc, "参考文献")
    old_refs = remove_following_paragraphs(ref_heading)
    if not old_refs:
        raise RuntimeError("未检测到原始参考文献段落")
    ref_template = old_refs[0]
    anchor = ref_heading
    for ref in references:
        anchor = insert_paragraph_after(anchor, ref, ref_template)

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
