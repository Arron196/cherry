from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v3.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v4_代码清单示例.docx"


@dataclass
class ListingSpec:
    match_prefix: str
    caption: str
    body: str


LISTINGS: list[ListingSpec] = [
    ListingSpec(
        match_prefix="SensorTask() {",
        caption="代码5-1 SensorTask任务伪代码",
        body="""SensorTask() {
    initSensors()
    cycle = 0
    while true:
        frame = {}
        frame.temp, frame.rh = SHT31_Read()
        frame.vibration_g = ADXL345_ReadMagnitude()
        if cycle >= 6:
            frame.co2_ppm = MHZ19B_ReadCO2()
        else:
            frame.co2_ppm = -1
        frame.timestamp = GetUTCTimestamp()
        frame.device_id = DEVICE_ID_CONST
        frame.batch_id = GetCurrentBatchId()
        xQueueSend(sensorQueue, frame, portMAX_DELAY)
        cycle += 1
        vTaskDelay(pdMS_TO_TICKS(30000))
}""",
    ),
    ListingSpec(
        match_prefix="SignTask() {",
        caption="代码5-2 SignTask任务伪代码",
        body="""SignTask() {
    while true:
        frame = xQueueReceive(sensorQueue, portMAX_DELAY)
        canonical = SerializeCanonical(frame)
        digest = HASH_SHA256(canonical)
        ATECC608A_Wakeup()
        signature_raw = ATECC608A_Sign(slot=0, digest)
        ATECC608A_Sleep()
        signature_der = EncodeDER(signature_raw)
        signed = BuildSignatureEnvelope(frame, signature_der, KEY_ID)
        xQueueSend(signedQueue, signed, portMAX_DELAY)
}""",
    ),
    ListingSpec(
        match_prefix="CommTask() {",
        caption="代码5-3 CommTask上传流程伪代码",
        body="""CommTask() {
    while true:
        signed = xQueueReceive(signedQueue, portMAX_DELAY)
        payload = SerializeJSON(signed)
        success = false
        if ESP8266_IsConnected():
            success = ESP8266_PostHTTPS(ENDPOINT, payload, IDEMPOTENCY_KEY)
        if not success:
            SX1278_TransmitFragmented(payload)
}""",
    ),
    ListingSpec(
        match_prefix="import json import hashlib",
        caption="代码5-4 Python端规范化哈希核心实现",
        body="""import json
import hashlib
from datetime import datetime, timezone

def _normalize_value(value):
    if isinstance(value, dict):
        return {k: _normalize_value(value[k]) for k in sorted(value.keys())}
    if isinstance(value, list):
        return [_normalize_value(item) for item in value]
    if isinstance(value, datetime):
        return value.astimezone(timezone.utc).strftime(
            '%Y-%m-%dT%H:%M:%S.%f+00:00'
        )
    if isinstance(value, str):
        return value.strip()
    return value

def compute_canonical_hash(data: dict) -> str:
    normalized = _normalize_value(data)
    canonical_json = json.dumps(
        normalized,
        ensure_ascii=True,
        separators=(',', ':'),
    )
    return hashlib.sha256(canonical_json.encode('utf-8')).hexdigest()""",
    ),
    ListingSpec(
        match_prefix="function normalizeValue(value: unknown): unknown {",
        caption="代码5-5 TypeScript端规范化哈希核心实现",
        body="""function normalizeValue(value: unknown): unknown {
  if (value instanceof Date) {
    return value.toISOString().replace('Z', '+00:00');
  }
  if (Array.isArray(value)) return value.map(normalizeValue);
  if (value !== null && typeof value === 'object') {
    return Object.fromEntries(
      Object.keys(value as object)
        .sort()
        .map(k => [k, normalizeValue((value as Record<string, unknown>)[k])])
    );
  }
  if (typeof value === 'string') return value.trim();
  return value;
}

async function computeCanonicalHash(data: object): Promise<string> {
  const normalized = normalizeValue(data);
  const json = JSON.stringify(normalized);
  const bytes = new TextEncoder().encode(json);
  const buf = await crypto.subtle.digest('SHA-256', bytes);
  return Array.from(new Uint8Array(buf))
    .map(b => b.toString(16).padStart(2, '0'))
    .join('');
}""",
    ),
    ListingSpec(
        match_prefix="try:",
        caption="代码5-6 canonical_hash冲突处理逻辑",
        body="""try:
    session.add(new_event)
    await session.flush()  # 触发 UNIQUE 约束检查
except IntegrityError:
    await session.rollback()
    existing = await session.execute(
        select(Event).where(Event.canonical_hash == canonical_hash)
    )
    return existing.scalar_one()""",
    ),
    ListingSpec(
        match_prefix="from cryptography.hazmat.primitives.asymmetric import ec",
        caption="代码5-7 ECDSA签名验证核心实现",
        body="""from cryptography.hazmat.primitives.asymmetric import ec
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.serialization import load_pem_public_key

def verify_ecdsa(public_key_pem: bytes, signature_der: bytes, canonical_hash: str) -> bool:
    try:
        public_key = load_pem_public_key(public_key_pem)
        message_bytes = bytes.fromhex(canonical_hash)
        public_key.verify(
            signature_der,
            message_bytes,
            ec.ECDSA(utils.Prehashed(hashes.SHA256())),
        )
        return True
    except InvalidSignature:
        return False""",
    ),
    ListingSpec(
        match_prefix="def verify_hmac(secret_key: bytes, canonical_hash: str,",
        caption="代码5-8 HMAC验证核心实现",
        body="""def verify_hmac(secret_key: bytes, canonical_hash: str, provided_mac: str) -> bool:
    expected = hmac.new(
        secret_key,
        canonical_hash.encode(),
        hashlib.sha256,
    ).hexdigest()
    return hmac.compare_digest(expected, provided_mac)""",
    ),
    ListingSpec(
        match_prefix="async def anchor_event(event: Event, adapter: AnchorAdapter):",
        caption="代码5-9 锚定任务核心流程",
        body="""async def anchor_event(event: Event, adapter: AnchorAdapter):
    event.ingest_status = IngestStatus.ANCHORING
    await db.flush()

    submission = AnchorSubmissionRecord(
        event_id=event.id,
        status=SubmissionStatus.PENDING,
    )
    db.add(submission)
    await db.flush()

    try:
        tx_hash = await adapter.anchor_event(event.canonical_hash, event.id)
        submission.transaction_hash = tx_hash
        receipt = await adapter.get_receipt(tx_hash)
        if await detect_reorg(receipt):
            submission.status = SubmissionStatus.REORGED
            event.ingest_status = IngestStatus.RECEIVED
        else:
            submission.status = SubmissionStatus.FINALIZED
            event.ingest_status = IngestStatus.ANCHORED
    except Exception:
        event.retry_count += 1
        event.ingest_status = (
            IngestStatus.DEAD_LETTER
            if event.retry_count >= 3 else IngestStatus.FAILED_RETRYING
        )""",
    ),
]


def ensure_rfonts(run, east_asia: str = "微软雅黑", ascii_font: str = "Consolas") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_paragraph_border(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        el = p_bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            p_bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "D6DCE8")


def set_paragraph_shading(paragraph: Paragraph, fill: str = "F5F7FA") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def style_code_paragraph(paragraph: Paragraph, body: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0.55)
    fmt.right_indent = Cm(0.35)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0

    run = paragraph.add_run(body)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(38, 52, 65)
    ensure_rfonts(run)
    set_paragraph_border(paragraph)
    set_paragraph_shading(paragraph)


def style_caption_paragraph(paragraph: Paragraph, caption: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(2)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0

    prefix, _, suffix = caption.partition(" ")
    run1 = paragraph.add_run(prefix)
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(10.5)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(28, 52, 84)
    ensure_rfonts(run1, east_asia="宋体", ascii_font="Times New Roman")

    run2 = paragraph.add_run(" " + suffix if suffix else "")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor(58, 74, 91)
    ensure_rfonts(run2, east_asia="宋体", ascii_font="Times New Roman")


def main() -> None:
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)
    paragraphs = list(doc.paragraphs)

    for spec in LISTINGS:
        target = None
        for para in paragraphs:
            text = para.text.strip().replace("\n", " ")
            if text.startswith(spec.match_prefix):
                target = para
                break
        if target is None:
            continue

        caption_para = insert_paragraph_before(target)
        style_caption_paragraph(caption_para, spec.caption)
        style_code_paragraph(target, spec.body)

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
