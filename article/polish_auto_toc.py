from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v14_自动目录修复2.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v15_自动目录精修.docx"


def set_rfonts(run, ascii_font: str = "Times New Roman", east_asia: str = "宋体") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def main() -> None:
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    heading = None
    heading_idx = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip().replace("\n", " ")
        if text in {"目  录", "目录", "目 录"}:
            heading = para
            heading_idx = idx
            break

    if heading is None or heading_idx is None:
        raise RuntimeError("未找到目录标题。")

    # 标题单独格式
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = heading.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    fmt.tab_stops.clear_all()
    for run in heading.runs:
        if run.text:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = "Times New Roman"
            set_rfonts(run, ascii_font="Times New Roman", east_asia="宋体")

    # 目录项格式
    for para in doc.paragraphs[heading_idx + 1 :]:
        text = para.text.replace("\n", " ")
        if not text.strip():
            continue
        if not re.search(r"\t\d+\s*$", text):
            break

        fmt = para.paragraph_format
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(20)
        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(Pt(467.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        clean = text.strip()
        if re.match(r"^\d+\.\d+\.\d+", clean):
            fmt.left_indent = Cm(0.85)
        elif re.match(r"^\d+\.\d+", clean):
            fmt.left_indent = Cm(0.42)
        else:
            fmt.left_indent = Cm(0)

        special = bool(re.match(r"^(摘要|关键词|Abstract|Keywords|参考文献|致谢|附录)", clean))
        for run in para.runs:
            if run.text:
                run.font.size = Pt(12)
                run.font.bold = special
                run.font.name = "Times New Roman"
                set_rfonts(run, ascii_font="Times New Roman", east_asia="宋体")

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
