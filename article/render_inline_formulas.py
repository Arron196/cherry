from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v4_代码清单示例.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v5_内联公式修复.docx"


def set_rfonts(run, ascii_font: str = "Cambria Math", east_asia: str = "Cambria Math") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_text(paragraph: Paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(12)


def add_math(paragraph: Paragraph, text: str, *, sub: bool = False, sup: bool = False, italic: bool | None = None) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.subscript = sub
    run.font.superscript = sup
    if italic is not None:
        run.font.italic = italic
    set_rfonts(run)


def rebuild_if_contains(doc: Document, needle: str, builder) -> bool:
    for para in doc.paragraphs:
        text = para.text.strip().replace("\n", " ")
        if needle in text:
            style = para.style
            align = para.alignment
            fmt = para.paragraph_format
            clear_paragraph(para)
            para.style = style
            para.alignment = align
            para.paragraph_format.left_indent = fmt.left_indent
            para.paragraph_format.right_indent = fmt.right_indent
            para.paragraph_format.first_line_indent = fmt.first_line_indent
            para.paragraph_format.line_spacing = fmt.line_spacing
            para.paragraph_format.line_spacing_rule = fmt.line_spacing_rule
            builder(para)
            return True
    return False


def main() -> None:
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    rebuild_if_contains(
        doc,
        "设第i个区块的哈希为H_i = SHA3-256(blockHeader_i)",
        lambda p: (
            add_text(p, "区块链（Blockchain）是一种以密码学哈希函数为纽带、将数据块按时间顺序链接而成的分布式账本结构。每个区块包含区块头（Block Header）和交易列表，区块头中存储前一区块的哈希值（parentHash）、交易Merkle树根（transactionsRoot）、时间戳、难度目标与Nonce等字段。设第i个区块的哈希为"),
            add_math(p, "H", italic=True),
            add_math(p, "i", sub=True, italic=True),
            add_math(p, " = SHA3-256(blockHeader", italic=False),
            add_math(p, "i", sub=True, italic=True),
            add_math(p, ")", italic=False),
            add_text(p, "，则链的完整性由递推关系 "),
            add_math(p, "blockHeader", italic=True),
            add_math(p, "i", sub=True, italic=True),
            add_math(p, ".parentHash = H", italic=False),
            add_math(p, "i-1", sub=True, italic=True),
            add_text(p, " 保证。若攻击者试图修改第i块中的某笔交易，则transactionsRoot随之变化，导致"),
            add_math(p, "H", italic=True),
            add_math(p, "i", sub=True, italic=True),
            add_text(p, "变化，进而使第i+1块的parentHash校验失败，整条链从第i+1块起全部无效。在权益证明（PoS）共识机制保护下，攻击者若要重写历史，需控制全网超过1/3的质押权益，攻击成本极高。"),
        ),
    )

    rebuild_if_contains(
        doc,
        "其中S_T和S_{RH}",
        lambda p: (
            add_text(p, "其中"),
            add_math(p, "S", italic=True),
            add_math(p, "T", sub=True, italic=True),
            add_text(p, "和"),
            add_math(p, "S", italic=True),
            add_math(p, "RH", sub=True, italic=True),
            add_text(p, "分别为16位温度和湿度原始值。"),
        ),
    )

    rebuild_if_contains(
        doc,
        "本系统采集三轴加速度的合矢量 a = sqrt",
        lambda p: (
            add_text(p, "在供应链物流监测场景中，ADXL345主要用于检测货物在装卸、运输过程中的振动冲击事件。系统配置ADXL345工作在±16g量程（分辨率3.9mg/LSB），通过I2C接口（与SHT31共用I2C1总线，但I2C地址不同：SHT31为0x44，ADXL345为0x53）实现数据读取。本系统采集三轴加速度的合矢量 "),
            add_math(p, "a", italic=True),
            add_math(p, " = √(", italic=False),
            add_math(p, "a", italic=True),
            add_math(p, "x", sub=True, italic=True),
            add_math(p, "2", sup=True, italic=False),
            add_math(p, " + ", italic=False),
            add_math(p, "a", italic=True),
            add_math(p, "y", sub=True, italic=True),
            add_math(p, "2", sup=True, italic=False),
            add_math(p, " + ", italic=False),
            add_math(p, "a", italic=True),
            add_math(p, "z", sub=True, italic=True),
            add_math(p, "2", sup=True, italic=False),
            add_math(p, ")", italic=False),
            add_text(p, " 减去重力分量1g，得到振动加速度值（单位：g），作为品质评分输入。"),
        ),
    )

    rebuild_if_contains(
        doc,
        "其中素数 p = 2^{256} - 2^{224}",
        lambda p: (
            add_text(p, "其中素数 "),
            add_math(p, "p", italic=True),
            add_math(p, " = 2", italic=False),
            add_math(p, "256", sup=True, italic=False),
            add_math(p, " - 2", italic=False),
            add_math(p, "224", sup=True, italic=False),
            add_math(p, " + 2", italic=False),
            add_math(p, "192", sup=True, italic=False),
            add_math(p, " + 2", italic=False),
            add_math(p, "96", sup=True, italic=False),
            add_math(p, " - 1", italic=False),
            add_text(p, "，参数 "),
            add_math(p, "a", italic=True),
            add_math(p, " = -3", italic=False),
            add_text(p, "，"),
            add_math(p, "b", italic=True),
            add_text(p, " 为特定常数。曲线上定义了“点加法”运算，选取基点 "),
            add_math(p, "G", italic=True),
            add_text(p, " 后，私钥 "),
            add_math(p, "d", italic=True),
            add_text(p, " 为随机大整数，公钥 "),
            add_math(p, "Q", italic=True),
            add_math(p, " = ", italic=False),
            add_math(p, "d", italic=True),
            add_math(p, " · ", italic=False),
            add_math(p, "G", italic=True),
            add_text(p, "（点乘运算）。ECDLP的困难性在于：已知 "),
            add_math(p, "G", italic=True),
            add_text(p, " 和 "),
            add_math(p, "Q", italic=True),
            add_text(p, "，求 "),
            add_math(p, "d", italic=True),
            add_text(p, " 在计算上不可行（无多项式时间算法），从而保证私钥无法从公钥逆推。"),
        ),
    )

    rebuild_if_contains(
        doc,
        "随机生成临时密钥 k ∈ [1, n-1]",
        lambda p: (
            add_text(p, "随机生成临时密钥 "),
            add_math(p, "k", italic=True),
            add_math(p, " ∈ [1, n-1]", italic=False),
            add_text(p, "（n为曲线阶）"),
        ),
    )

    rebuild_if_contains(
        doc,
        "计算椭圆曲线点 (x_1, y_1) = k · G",
        lambda p: (
            add_text(p, "计算椭圆曲线点 "),
            add_math(p, "(", italic=False),
            add_math(p, "x", italic=True),
            add_math(p, "1", sub=True, italic=False),
            add_math(p, ", ", italic=False),
            add_math(p, "y", italic=True),
            add_math(p, "1", sub=True, italic=False),
            add_math(p, ")", italic=False),
            add_math(p, " = ", italic=False),
            add_math(p, "k", italic=True),
            add_math(p, " · ", italic=False),
            add_math(p, "G", italic=True),
            add_text(p, "，令 "),
            add_math(p, "r", italic=True),
            add_math(p, " = ", italic=False),
            add_math(p, "x", italic=True),
            add_math(p, "1", sub=True, italic=False),
            add_math(p, " mod n", italic=False),
        ),
    )

    rebuild_if_contains(
        doc,
        "计算 s = k^{-1}(e + r · d) mod n",
        lambda p: (
            add_text(p, "计算 "),
            add_math(p, "s", italic=True),
            add_math(p, " = ", italic=False),
            add_math(p, "k", italic=True),
            add_math(p, "-1", sup=True, italic=False),
            add_math(p, "(", italic=False),
            add_math(p, "e", italic=True),
            add_math(p, " + ", italic=False),
            add_math(p, "r", italic=True),
            add_math(p, " · ", italic=False),
            add_math(p, "d", italic=True),
            add_math(p, ") mod n", italic=False),
        ),
    )

    rebuild_if_contains(
        doc,
        "验证方持有公钥 Q 时，可独立验证：计算 u_1 = e · s^{-1} mod n",
        lambda p: (
            add_text(p, "验证方持有公钥 "),
            add_math(p, "Q", italic=True),
            add_text(p, " 时，可独立验证：计算 "),
            add_math(p, "u", italic=True),
            add_math(p, "1", sub=True, italic=False),
            add_math(p, " = ", italic=False),
            add_math(p, "e", italic=True),
            add_math(p, " · ", italic=False),
            add_math(p, "s", italic=True),
            add_math(p, "-1", sup=True, italic=False),
            add_math(p, " mod n", italic=False),
            add_text(p, "，"),
            add_math(p, "u", italic=True),
            add_math(p, "2", sub=True, italic=False),
            add_math(p, " = ", italic=False),
            add_math(p, "r", italic=True),
            add_math(p, " · ", italic=False),
            add_math(p, "s", italic=True),
            add_math(p, "-1", sup=True, italic=False),
            add_math(p, " mod n", italic=False),
            add_text(p, "，点 "),
            add_math(p, "(x, y)", italic=False),
            add_math(p, " = ", italic=False),
            add_math(p, "u", italic=True),
            add_math(p, "1", sub=True, italic=False),
            add_math(p, "G + ", italic=False),
            add_math(p, "u", italic=True),
            add_math(p, "2", sub=True, italic=False),
            add_math(p, "Q", italic=False),
            add_text(p, "，若 "),
            add_math(p, "x", italic=True),
            add_math(p, " ≡ ", italic=False),
            add_math(p, "r", italic=True),
            add_math(p, " (mod n)", italic=False),
            add_text(p, " 则签名有效。"),
        ),
    )

    rebuild_if_contains(
        doc,
        "其中B_i为第i个指标的分段得分",
        lambda p: (
            add_text(p, "其中"),
            add_math(p, "B", italic=True),
            add_math(p, "i", sub=True, italic=False),
            add_text(p, "为第i个指标的分段得分（BandScore），"),
            add_math(p, "w", italic=True),
            add_math(p, "i", sub=True, italic=False),
            add_text(p, "为权重。各指标的分段规则与权重由rules.yml配置文件定义，支持热重载（无需重启服务），便于业务方根据不同品种、不同季节调整阈值："),
        ),
    )

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
