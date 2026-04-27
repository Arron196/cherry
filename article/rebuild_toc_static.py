from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
INPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v6_目录修复.docx"
OUTPUT_DOC = ROOT / "论文——模板编辑_按规范排版_配图版_v12_目录静态重建.docx"


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_rfonts(run, ascii_font: str = "Times New Roman", east_asia: str = "宋体") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def style_heading(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    run = paragraph.add_run("目  录")
    run.font.size = Pt(16)
    run.font.bold = True
    set_rfonts(run, ascii_font="Times New Roman", east_asia="宋体")


def style_blank(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0


def style_toc_entry(paragraph: Paragraph, title: str, page: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    fmt.tab_stops.clear_all()
    fmt.tab_stops.add_tab_stop(Pt(467.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    if re.match(r"^\d+\.\d+\.\d+", title):
        fmt.left_indent = Cm(0.85)
    elif re.match(r"^\d+\.\d+", title):
        fmt.left_indent = Cm(0.42)
    else:
        fmt.left_indent = Cm(0)

    special_bold = bool(re.match(r"^(摘要|关键词|Abstract|Keywords|参考文献|致谢|附录)$", title))

    run = paragraph.add_run(f"{title}\t{page}")
    run.font.size = Pt(12)
    run.font.bold = special_bold
    set_rfonts(run, ascii_font="Times New Roman", east_asia="宋体")


def main() -> None:
    shutil.copy2(INPUT_DOC, OUTPUT_DOC)
    doc = Document(OUTPUT_DOC)

    # locate toc title / entries / section break paragraph
    heading_idx = None
    section_break_idx = None
    entries: list[tuple[str, str]] = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip().replace("\n", " ")
        if heading_idx is None and text in {"目  录", "目录", "目 录"}:
            heading_idx = idx
            continue

        if heading_idx is not None and section_break_idx is None:
            ppr = para._p.pPr
            if ppr is not None and ppr.sectPr is not None:
                section_break_idx = idx
                break
            m = re.match(r"^(.*)\t(\d+)\s*$", para.text.replace("\n", " "))
            if m:
                entries.append((m.group(1), m.group(2)))

    if heading_idx is None or section_break_idx is None or not entries:
        raise RuntimeError("未能识别目录区域。")

    # delete old toc content between heading and section break, keep heading and section break
    for idx in range(section_break_idx - 1, heading_idx, -1):
        remove_paragraph(doc.paragraphs[idx])

    heading = doc.paragraphs[heading_idx]
    style_heading(heading)

    section_break_para = doc.paragraphs[heading_idx + 1]

    # blank line
    blank = insert_paragraph_before(section_break_para)
    style_blank(blank)

    # insert entries in natural order before section break
    for title, page in entries:
        para = insert_paragraph_before(section_break_para)
        style_toc_entry(para, title, page)

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
