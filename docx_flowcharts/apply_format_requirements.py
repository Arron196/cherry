from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INPUT_DOCX = Path(r"C:\Users\benja\Dropbox\毕业论文\任相荣论文1_流程图版.docx")
OUTPUT_DOCX = Path(r"C:\Users\benja\Dropbox\毕业论文\任相荣论文1_按规范格式调整版.docx")

BODY_START = "绪论"
BACK_MATTER = {"参考文献", "致 谢", "致谢", "附录"}


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: Pt | None = None, bold=None) -> None:
    run.font.name = latin
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_style_font(style, east_asia: str, latin: str, size_pt: float, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def set_outline_level(style, level: int) -> None:
    pPr = style._element.get_or_add_pPr()
    num_pr = pPr.find(qn("w:numPr"))
    if num_pr is not None:
        pPr.remove(num_pr)
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_toc_field(paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])
    set_run_font(run, size=Pt(12), bold=False)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=Pt(9))


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def ensure_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(paragraph)
    add_page_field(paragraph)
    sectPr = section._sectPr
    pg_num = sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sectPr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
    if len(doc.sections) > 0:
        # Front matter in the provided template has no visible page number.
        doc.sections[0].footer.is_linked_to_previous = False
        for p in doc.sections[0].footer.paragraphs:
            clear_paragraph(p)
    if len(doc.sections) > 1:
        ensure_page_number(doc.sections[1])


def configure_styles(doc: Document) -> None:
    set_style_font(doc.styles["Normal"], "宋体", "Times New Roman", 12, False)
    normal_pf = doc.styles["Normal"].paragraph_format
    normal_pf.line_spacing = 1.5
    normal_pf.first_line_indent = Cm(0.74)
    normal_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(0)

    set_style_font(doc.styles["论文一级标题"], "宋体", "Times New Roman", 15, True)
    set_outline_level(doc.styles["论文一级标题"], 0)
    doc.styles["论文一级标题"].paragraph_format.line_spacing = 1.5
    doc.styles["论文一级标题"].paragraph_format.first_line_indent = None
    doc.styles["论文一级标题"].paragraph_format.space_before = Pt(0)
    doc.styles["论文一级标题"].paragraph_format.space_after = Pt(0)

    set_style_font(doc.styles["论文二级标题"], "宋体", "Times New Roman", 14, True)
    set_outline_level(doc.styles["论文二级标题"], 1)
    doc.styles["论文二级标题"].paragraph_format.line_spacing = 1.5
    doc.styles["论文二级标题"].paragraph_format.first_line_indent = None
    doc.styles["论文二级标题"].paragraph_format.space_before = Pt(0)
    doc.styles["论文二级标题"].paragraph_format.space_after = Pt(0)

    set_style_font(doc.styles["论文三级标题"], "宋体", "Times New Roman", 12, True)
    set_outline_level(doc.styles["论文三级标题"], 2)
    doc.styles["论文三级标题"].paragraph_format.line_spacing = 1.5
    doc.styles["论文三级标题"].paragraph_format.first_line_indent = None
    doc.styles["论文三级标题"].paragraph_format.space_before = Pt(0)
    doc.styles["论文三级标题"].paragraph_format.space_after = Pt(0)

    set_style_font(doc.styles["表题"], "宋体", "Times New Roman", 10.5, False)
    doc.styles["表题"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["表题"].paragraph_format.line_spacing = 1.5
    doc.styles["表题"].paragraph_format.first_line_indent = None
    doc.styles["表题"].paragraph_format.space_before = Pt(0)
    doc.styles["表题"].paragraph_format.space_after = Pt(0)

    set_style_font(doc.styles["参考文献格式"], "宋体", "Times New Roman", 10.5, False)
    refs_pf = doc.styles["参考文献格式"].paragraph_format
    refs_pf.line_spacing = 1.5
    refs_pf.left_indent = Cm(0.74)
    refs_pf.first_line_indent = Cm(-0.74)
    refs_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    refs_pf.space_before = Pt(0)
    refs_pf.space_after = Pt(0)

    for toc_style_name, indent_cm in [("TOC 1", 0), ("TOC 2", 0.37), ("TOC 3", 0.74)]:
        if toc_style_name in doc.styles:
            set_style_font(doc.styles[toc_style_name], "宋体", "Times New Roman", 12, False)
            pf = doc.styles[toc_style_name].paragraph_format
            pf.line_spacing = 1.5
            pf.left_indent = Cm(indent_cm)
            pf.first_line_indent = None
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)


def insert_toc(doc: Document) -> None:
    if any(p.text.strip() == "目  录" for p in doc.paragraphs):
        return
    marker = None
    for paragraph in doc.paragraphs:
        if paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None:
            marker = paragraph
            break
    if marker is None:
        marker = doc.paragraphs[12]

    before_break = marker.insert_paragraph_before("")
    before_break.paragraph_format.first_line_indent = None
    before_break.add_run().add_break(WD_BREAK.PAGE)

    title = marker.insert_paragraph_before("目  录")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.line_spacing = 1.5
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    for run in title.runs:
        set_run_font(run, size=Pt(16), bold=True)

    toc = marker.insert_paragraph_before("")
    add_toc_field(toc)
    spacer = marker.insert_paragraph_before("")
    spacer.paragraph_format.first_line_indent = None
    after_break = marker.insert_paragraph_before("")
    after_break.paragraph_format.first_line_indent = None
    after_break.add_run().add_break(WD_BREAK.PAGE)


def replace_paragraph_text(paragraph, text: str) -> None:
    old_runs = list(paragraph.runs)
    if old_runs:
        old_runs[0].text = text
        for run in old_runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def strip_existing_heading_number(text: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*\s+", "", text).strip()


def number_headings(doc: Document) -> None:
    chapter = 0
    second = 0
    third = 0
    in_body = False
    in_back = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == BODY_START:
            in_body = True
        if paragraph.style.name == "论文一级标题" and text in BACK_MATTER:
            in_back = True
        if not in_body or in_back or not text:
            continue
        if paragraph.style.name == "论文一级标题":
            chapter += 1
            second = 0
            third = 0
            replace_paragraph_text(paragraph, f"{chapter} {strip_existing_heading_number(text)}")
        elif paragraph.style.name == "论文二级标题":
            second += 1
            third = 0
            replace_paragraph_text(paragraph, f"{chapter}.{second} {strip_existing_heading_number(text)}")
        elif paragraph.style.name == "论文三级标题":
            third += 1
            replace_paragraph_text(paragraph, f"{chapter}.{second}.{third} {strip_existing_heading_number(text)}")


def caption_text_name(text: str) -> str:
    text = text.strip()
    if text.startswith("流程图"):
        return re.sub(r"^流程图\s*\d+(?:[-－]\d+)?\s*", "", text)
    return re.sub(r"^图\s*\d+\s*", "", text)


def renumber_figures(doc: Document) -> None:
    fig_no = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^(图\s*\d+|流程图\s*\d+(?:[-－]\d+)?)", text):
            fig_no += 1
            replace_paragraph_text(paragraph, f"图{fig_no} {caption_text_name(text)}")


def normalize_paragraphs(doc: Document) -> None:
    in_refs = False
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name
        pf = paragraph.paragraph_format

        if text == "参考文献":
            in_refs = True
        if text in {"致 谢", "致谢", "附录"}:
            in_refs = False

        if idx <= 24:
            # Front matter follows the school template: centered title block,
            # justified abstract, no first-line indent for keywords.
            if text.startswith(("关键词", "Keywords")):
                pf.first_line_indent = None
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif text.startswith(("摘", "Abstract")):
                pf.first_line_indent = Cm(0.74)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                if idx == 12:
                    set_run_font(run, size=Pt(16), bold=True)
                elif idx == 19:
                    set_run_font(run, size=Pt(14), bold=True)
                elif text:
                    set_run_font(run, size=Pt(12))
            continue

        if style in {"论文一级标题", "论文二级标题", "论文三级标题"}:
            pf.first_line_indent = None
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif re.match(r"^图\d+\s", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=Pt(10.5), bold=False)
        elif re.match(r"^表\d+\s", text):
            paragraph.style = doc.styles["表题"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            for run in paragraph.runs:
                set_run_font(run, size=Pt(10.5), bold=False)
        elif in_refs and re.match(r"^\[\d+\]", text):
            paragraph.style = doc.styles["参考文献格式"]
            for run in paragraph.runs:
                set_run_font(run, size=Pt(10.5), bold=False)
        elif style == "Normal" and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing = 1.5
            pf.first_line_indent = Cm(0.74)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

        for run in paragraph.runs:
            if style == "论文一级标题":
                set_run_font(run, size=Pt(15), bold=True)
            elif style == "论文二级标题":
                set_run_font(run, size=Pt(14), bold=True)
            elif style == "论文三级标题":
                set_run_font(run, size=Pt(12), bold=True)
            elif style == "参考文献格式":
                set_run_font(run, size=Pt(10.5), bold=False)
            elif re.match(r"^图\d+\s", text):
                set_run_font(run, size=Pt(10.5), bold=False)
            elif idx > 24 and text:
                set_run_font(run, size=Pt(12))


def resize_inserted_flowcharts(doc: Document) -> None:
    # Flowchart raster images are inserted around the captions that used to be
    # code blocks. Keep them inside the 16 cm text width with breathing room.
    caption_indexes = {
        i
        for i, p in enumerate(doc.paragraphs)
        if re.match(r"^图\d+\s", p.text.strip()) and any(key in p.text for key in [
            "SensorTask任务流程图",
            "SignTask任务流程图",
            "CommTask上传流程图",
            "规范化哈希流程图",
            "canonical_hash冲突处理流程图",
            "ECDSA签名验证流程图",
            "HMAC验证流程图",
            "锚定任务核心流程图",
        ])
    }
    for idx in caption_indexes:
        if idx == 0:
            continue
        pic_para = doc.paragraphs[idx - 1]
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.paragraph_format.first_line_indent = None
        for shape in pic_para._p.xpath(".//wp:inline"):
            extent = shape.xpath("./wp:extent")[0]
            cx = int(extent.get("cx"))
            cy = int(extent.get("cy"))
            target_cx = Cm(14.2)
            ratio = target_cx / cx
            extent.set("cx", str(int(target_cx)))
            extent.set("cy", str(int(cy * ratio)))
            for ext in shape.xpath(".//a:ext"):
                if ext.get("cx") == str(cx):
                    ext.set("cx", str(int(target_cx)))
                    ext.set("cy", str(int(cy * ratio)))


def main() -> None:
    doc = Document(INPUT_DOCX)
    set_margins(doc)
    configure_styles(doc)
    number_headings(doc)
    renumber_figures(doc)
    insert_toc(doc)
    normalize_paragraphs(doc)
    resize_inserted_flowcharts(doc)
    doc.save(OUTPUT_DOCX)
    print(f"saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
