from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


INPUT_DOCX = Path(r"C:\Users\benja\Dropbox\毕业论文\任相荣论文1.docx")
OUTPUT_DOCX = Path(r"C:\Users\benja\Dropbox\毕业论文\任相荣论文1_流程图版.docx")
OUT_DIR = Path(r"C:\学校\cherry\docx_flowcharts\generated")

FONT_REGULAR = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"


PALETTE = {
    "bg": "#ffffff",
    "line": "#45515f",
    "text": "#172033",
    "start": "#dff4e8",
    "process": "#eaf2ff",
    "queue": "#efeaff",
    "decision": "#fff3c7",
    "error": "#ffe5e5",
    "stroke": "#7d8da3",
    "green": "#5aa47b",
    "blue": "#6b8fcf",
    "purple": "#8e78bf",
    "yellow": "#c7a646",
    "red": "#d26b6b",
}


@dataclass
class Node:
    x: int
    y: int
    w: int
    h: int
    shape: str
    text: str

    @property
    def left(self) -> int:
        return self.x

    @property
    def right(self) -> int:
        return self.x + self.w

    @property
    def top(self) -> int:
        return self.y

    @property
    def bottom(self) -> int:
        return self.y + self.h

    @property
    def cx(self) -> int:
        return self.x + self.w // 2

    @property
    def cy(self) -> int:
        return self.y + self.h // 2


class FlowChart:
    def __init__(self, width: int = 1800) -> None:
        self.width = width
        self.height = 3200
        self.margin_x = 100
        self.image = Image.new("RGB", (self.width, self.height), PALETTE["bg"])
        self.draw = ImageDraw.Draw(self.image)
        self.font = ImageFont.truetype(FONT_REGULAR, 42)
        self.small_font = ImageFont.truetype(FONT_REGULAR, 34)
        self.bold_font = ImageFont.truetype(FONT_BOLD, 42)
        self.max_y = 0

    def _wrap(self, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
        lines: list[str] = []
        current = ""
        for char in text:
            if char == "\n":
                if current:
                    lines.append(current)
                current = ""
                continue
            trial = current + char
            if self.draw.textlength(trial, font=font) <= max_width or not current:
                current = trial
            else:
                lines.append(current.rstrip())
                current = char.lstrip()
        if current:
            lines.append(current.rstrip())
        return lines

    def text_size(self, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> tuple[int, int, list[str]]:
        lines = self._wrap(text, font, max_width)
        widths = [self.draw.textlength(line, font=font) for line in lines] or [0]
        bbox = self.draw.textbbox((0, 0), "测", font=font)
        line_h = bbox[3] - bbox[1] + 14
        return int(max(widths)), line_h * len(lines), lines

    def node(self, x: int, y: int, w: int, h: int, text: str, shape: str = "process") -> Node:
        n = Node(x, y, w, h, shape, text)
        fill = {
            "start": PALETTE["start"],
            "process": PALETTE["process"],
            "queue": PALETTE["queue"],
            "decision": PALETTE["decision"],
            "error": PALETTE["error"],
        }.get(shape, PALETTE["process"])
        outline = {
            "start": PALETTE["green"],
            "process": PALETTE["blue"],
            "queue": PALETTE["purple"],
            "decision": PALETTE["yellow"],
            "error": PALETTE["red"],
        }.get(shape, PALETTE["blue"])
        if shape == "decision":
            pts = [(n.cx, n.top), (n.right, n.cy), (n.cx, n.bottom), (n.left, n.cy)]
            self.draw.polygon(pts, fill=fill, outline=outline)
            self.draw.line(pts + [pts[0]], fill=outline, width=4, joint="curve")
            text_max = int(w * 0.62)
        elif shape == "start":
            self.draw.rounded_rectangle([n.left, n.top, n.right, n.bottom], radius=h // 2, fill=fill, outline=outline, width=4)
            text_max = w - 70
        else:
            self.draw.rounded_rectangle([n.left, n.top, n.right, n.bottom], radius=22, fill=fill, outline=outline, width=4)
            text_max = w - 70
        _, text_h, lines = self.text_size(text, self.font, text_max)
        bbox = self.draw.textbbox((0, 0), "测", font=self.font)
        line_h = bbox[3] - bbox[1] + 14
        ty = n.cy - text_h // 2 - 3
        for line in lines:
            tw = self.draw.textlength(line, font=self.font)
            self.draw.text((n.cx - tw / 2, ty), line, font=self.font, fill=PALETTE["text"])
            ty += line_h
        self.max_y = max(self.max_y, n.bottom)
        return n

    def arrow(self, points: Iterable[tuple[int, int]], label: str | None = None, label_offset: tuple[int, int] = (0, 0)) -> None:
        pts = list(points)
        self.draw.line(pts, fill=PALETTE["line"], width=5, joint="curve")
        if len(pts) >= 2:
            x1, y1 = pts[-2]
            x2, y2 = pts[-1]
            self._arrowhead(x1, y1, x2, y2)
        if label:
            mid = pts[len(pts) // 2]
            tx = mid[0] + label_offset[0]
            ty = mid[1] + label_offset[1]
            pad_x, pad_y = 16, 8
            bbox = self.draw.textbbox((tx, ty), label, font=self.small_font)
            self.draw.rounded_rectangle(
                [bbox[0] - pad_x, bbox[1] - pad_y, bbox[2] + pad_x, bbox[3] + pad_y],
                radius=12,
                fill=PALETTE["bg"],
                outline="#d6dde8",
                width=2,
            )
            self.draw.text((tx, ty), label, font=self.small_font, fill=PALETTE["text"])

    def _arrowhead(self, x1: int, y1: int, x2: int, y2: int) -> None:
        import math

        angle = math.atan2(y2 - y1, x2 - x1)
        length = 22
        spread = math.radians(26)
        p1 = (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread))
        p2 = (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread))
        self.draw.polygon([(x2, y2), p1, p2], fill=PALETTE["line"])

    def down(self, a: Node, b: Node, label: str | None = None) -> None:
        self.arrow([(a.cx, a.bottom), (b.cx, b.top)], label=label, label_offset=(20, -22))

    def save(self, path: Path) -> None:
        crop_h = min(self.height, self.max_y + 110)
        cropped = self.image.crop((0, 0, self.width, crop_h))
        cropped.save(path, dpi=(300, 300))


def make_simple_vertical(path: Path, steps: list[tuple[str, str]], box_w: int = 1020, box_h: int = 118, gap: int = 58) -> None:
    chart = FlowChart()
    x = (chart.width - box_w) // 2
    nodes: list[Node] = []
    y = 70
    for text, shape in steps:
        h = box_h + (36 if len(text) > 28 else 0)
        nodes.append(chart.node(x, y, box_w, h, text, shape))
        y += h + gap
    for a, b in zip(nodes, nodes[1:]):
        chart.down(a, b)
    chart.save(path)


def sensor_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h, "开始：SensorTask", "start"); y += h + gap
    init = c.node(x, y, w, h, "初始化传感器，cycle = 0", "process"); y += h + gap
    loop = c.node(x, y, w, h, "进入30秒采集周期", "queue"); y += h + gap
    read1 = c.node(x, y, w, h + 18, "读取SHT31温湿度与ADXL345振动幅值", "process"); y += h + gap + 18
    decision = c.node(x + 190, y, 640, 210, "cycle >= 6 ?", "decision"); y += 210 + gap
    left = c.node(170, y, 620, h + 20, "读取MH-Z19B CO2浓度", "process")
    right = c.node(1010, y, 620, h + 20, "CO2字段置为-1", "process")
    y += h + 95
    meta = c.node(x, y, w, h + 28, "补充timestamp、device_id与batch_id", "process"); y += h + gap + 28
    send = c.node(x, y, w, h, "发送数据帧到sensorQueue", "queue"); y += h + gap
    delay = c.node(x, y, w, h + 18, "cycle加1，延时30秒后进入下一轮", "process")
    c.down(start, init); c.down(init, loop); c.down(loop, read1); c.down(read1, decision)
    c.arrow([(decision.left + 55, decision.cy), (left.cx, decision.cy), (left.cx, left.top)], "是", (-80, -44))
    c.arrow([(decision.right - 55, decision.cy), (right.cx, decision.cy), (right.cx, right.top)], "否", (50, -44))
    c.arrow([(left.cx, left.bottom), (left.cx, meta.top - 25), (meta.cx, meta.top - 25), (meta.cx, meta.top)])
    c.arrow([(right.cx, right.bottom), (right.cx, meta.top - 25), (meta.cx, meta.top - 25), (meta.cx, meta.top)])
    c.down(meta, send); c.down(send, delay)
    c.arrow([(delay.right, delay.cy), (1665, delay.cy), (1665, loop.cy), (loop.right, loop.cy)], "循环", (18, -44))
    c.save(path)


def comm_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h, "开始：CommTask", "start"); y += h + gap
    wait = c.node(x, y, w, h, "等待signedQueue中的签名数据", "queue"); y += h + gap
    json = c.node(x, y, w, h, "序列化为JSON，success = false", "process"); y += h + gap
    wifi = c.node(x + 190, y, 640, 210, "Wi-Fi已连接?", "decision"); y += 210 + gap
    post = c.node(165, y, 650, h + 20, "通过ESP8266执行HTTPS POST", "process")
    lora = c.node(985, y, 650, h + 20, "使用SX1278进行LoRa分片发送", "process")
    y += h + 95
    ok = c.node(165, y, 650, 190, "POST成功?", "decision")
    done = c.node(985, y + 35, 650, h, "完成本次上传并等待下一条", "start")
    c.down(start, wait); c.down(wait, json); c.down(json, wifi)
    c.arrow([(wifi.left + 55, wifi.cy), (post.cx, wifi.cy), (post.cx, post.top)], "是", (-80, -44))
    c.arrow([(wifi.right - 55, wifi.cy), (lora.cx, wifi.cy), (lora.cx, lora.top)], "否", (45, -44))
    c.down(post, ok)
    c.arrow([(ok.right - 55, ok.cy), (done.left, done.cy)], "是", (25, -44))
    c.arrow([(ok.cx, ok.bottom), (ok.cx, lora.bottom + 55), (lora.cx, lora.bottom + 55), (lora.cx, lora.bottom)], "否", (20, 10))
    c.arrow([(lora.right, lora.cy), (1690, lora.cy), (1690, done.cy), (done.right, done.cy)])
    c.arrow([(done.right, done.cy), (1710, done.cy), (1710, wait.cy), (wait.right, wait.cy)], "循环", (10, -44))
    c.save(path)


def canonical_branch_chart(path: Path, title_text: str, hash_step: str) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h, title_text, "start"); y += h + gap
    walk = c.node(x, y, w, h, "递归遍历输入数据", "queue"); y += h + gap
    decision = c.node(x + 190, y, 640, 210, "当前值类型?", "decision"); y += 245
    b1 = c.node(105, y, 470, h + 65, "Object / Dict：按键排序后递归处理", "process")
    b2 = c.node(665, y, 470, h + 65, "Array / List：逐项递归处理", "process")
    b3 = c.node(1225, y, 470, h + 65, "Date / String：UTC格式化或trim", "process")
    y += h + 130
    merge = c.node(x, y, w, h, "得到normalized对象", "queue"); y += h + gap
    json = c.node(x, y, w, h + 18, "生成紧凑JSON字节串", "process"); y += h + gap + 18
    digest = c.node(x, y, w, h + 28, hash_step, "process"); y += h + gap + 28
    out = c.node(x, y, w, h, "输出64位小写十六进制哈希", "start")
    c.down(start, walk); c.down(walk, decision)
    c.arrow([(decision.left + 45, decision.cy), (b1.cx, decision.cy), (b1.cx, b1.top)], "分支", (-118, -44))
    c.arrow([(decision.cx, decision.bottom), (b2.cx, b2.top)])
    c.arrow([(decision.right - 45, decision.cy), (b3.cx, decision.cy), (b3.cx, b3.top)], "分支", (35, -44))
    for b in (b1, b2, b3):
        c.arrow([(b.cx, b.bottom), (b.cx, merge.top - 25), (merge.cx, merge.top - 25), (merge.cx, merge.top)])
    c.down(merge, json); c.down(json, digest); c.down(digest, out)
    c.save(path)


def hmac_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h, "输入secret_key、canonical_hash\n与provided_mac", "start"); y += h + gap
    calc = c.node(x, y, w, h + 24, "计算HMAC-SHA256得到expected", "process"); y += h + gap + 24
    cmpn = c.node(x + 190, y, 640, 210, "compare_digest一致?", "decision"); y += 210 + gap
    true = c.node(175, y, 610, h, "返回True", "start")
    false = c.node(1015, y, 610, h, "返回False", "error")
    c.down(start, calc); c.down(calc, cmpn)
    c.arrow([(cmpn.left + 55, cmpn.cy), (true.cx, cmpn.cy), (true.cx, true.top)], "是", (-80, -44))
    c.arrow([(cmpn.right - 55, cmpn.cy), (false.cx, cmpn.cy), (false.cx, false.top)], "否", (45, -44))
    c.save(path)


def ecdsa_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h + 22, "输入公钥、DER签名与canonical_hash", "start"); y += h + gap + 22
    load = c.node(x, y, w, h, "加载PEM公钥", "process"); y += h + gap
    bytesn = c.node(x, y, w, h, "canonical_hash十六进制转bytes", "process"); y += h + gap
    verify = c.node(x, y, w, h + 32, "使用ECDSA + Prehashed(SHA-256)执行验证", "process"); y += h + gap + 32
    dec = c.node(x + 190, y, 640, 210, "验证是否抛出\nInvalidSignature?", "decision"); y += 210 + gap
    true = c.node(175, y, 610, h, "未抛异常：返回True", "start")
    false = c.node(1015, y, 610, h, "捕获异常：返回False", "error")
    c.down(start, load); c.down(load, bytesn); c.down(bytesn, verify); c.down(verify, dec)
    c.arrow([(dec.left + 55, dec.cy), (true.cx, dec.cy), (true.cx, true.top)], "否", (-80, -44))
    c.arrow([(dec.right - 55, dec.cy), (false.cx, dec.cy), (false.cx, false.top)], "是", (45, -44))
    c.save(path)


def conflict_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 118, 55
    y = 60
    start = c.node(x, y, w, h, "新TraceEvent准备入库", "start"); y += h + gap
    add = c.node(x, y, w, h, "session.add(new_event)", "process"); y += h + gap
    flush = c.node(x, y, w, h, "flush触发UNIQUE约束检查", "process"); y += h + gap
    dec = c.node(x + 190, y, 640, 210, "canonical_hash冲突?", "decision"); y += 210 + gap
    ok = c.node(155, y, 650, h, "无冲突：写入成功", "start")
    rollback = c.node(995, y, 650, h, "有冲突：rollback", "error")
    y += h + gap
    query = c.node(995, y, 650, h + 35, "按canonical_hash查询已有事件", "process"); y += h + gap + 35
    ret = c.node(995, y, 650, h, "返回已有事件", "start")
    c.down(start, add); c.down(add, flush); c.down(flush, dec)
    c.arrow([(dec.left + 55, dec.cy), (ok.cx, dec.cy), (ok.cx, ok.top)], "否", (-80, -44))
    c.arrow([(dec.right - 55, dec.cy), (rollback.cx, dec.cy), (rollback.cx, rollback.top)], "是", (45, -44))
    c.down(rollback, query); c.down(query, ret)
    c.save(path)


def anchor_chart(path: Path) -> None:
    c = FlowChart()
    x, w, h, gap = 390, 1020, 112, 48
    y = 55
    start = c.node(x, y, w, h, "输入待锚定Event", "start"); y += h + gap
    anchoring = c.node(x, y, w, h, "event状态置为ANCHORING并flush", "process"); y += h + gap
    submission = c.node(x, y, w, h + 18, "创建PENDING提交记录并flush", "process"); y += h + gap + 18
    tx = c.node(x, y, w, h, "调用adapter.anchor_event获得tx_hash", "process"); y += h + gap
    receipt = c.node(x, y, w, h, "查询交易receipt", "process"); y += h + gap
    reorg = c.node(x + 190, y, 640, 205, "检测到链重组?", "decision"); y += 205 + gap
    yes = c.node(140, y, 700, h + 18, "submission=REORGED，event回到RECEIVED", "error")
    no = c.node(960, y, 700, h + 18, "submission=FINALIZED，event置为ANCHORED", "start")
    y += h + 95
    exc = c.node(x + 190, y, 640, 205, "锚定过程异常?", "decision"); y += 205 + gap
    retry = c.node(x, y, w, h, "retry_count加1", "process"); y += h + gap
    retry_dec = c.node(x + 190, y, 640, 205, "retry_count >= 3?", "decision"); y += 205 + gap
    retrying = c.node(140, y, 700, h, "置为FAILED_RETRYING", "process")
    dead = c.node(960, y, 700, h, "置为DEAD_LETTER", "error")
    c.down(start, anchoring); c.down(anchoring, submission); c.down(submission, tx); c.down(tx, receipt); c.down(receipt, reorg)
    c.arrow([(reorg.left + 55, reorg.cy), (yes.cx, reorg.cy), (yes.cx, yes.top)], "是", (-80, -44))
    c.arrow([(reorg.right - 55, reorg.cy), (no.cx, reorg.cy), (no.cx, no.top)], "否", (45, -44))
    c.arrow([(tx.right, tx.cy), (1665, tx.cy), (1665, exc.cy), (exc.right, exc.cy)], "异常", (15, -44))
    c.down(exc, retry, "是")
    c.down(retry, retry_dec)
    c.arrow([(retry_dec.left + 55, retry_dec.cy), (retrying.cx, retry_dec.cy), (retrying.cx, retrying.top)], "否", (-80, -44))
    c.arrow([(retry_dec.right - 55, retry_dec.cy), (dead.cx, retry_dec.cy), (dead.cx, dead.top)], "是", (45, -44))
    c.save(path)


def generate_images() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images: dict[str, Path] = {}

    specs = [
        ("代码5-1 SensorTask任务伪代码", "流程图5-1 SensorTask任务流程图", "flow_5_1_sensortask.png", sensor_chart),
        (
            "代码5-2 SignTask任务伪代码",
            "流程图5-2 SignTask任务流程图",
            "flow_5_2_signtask.png",
            lambda p: make_simple_vertical(
                p,
                [
                    ("开始：SignTask", "start"),
                    ("等待sensorQueue中的数据帧", "queue"),
                    ("规范化序列化数据帧", "process"),
                    ("使用HASH硬件计算SHA-256摘要", "process"),
                    ("唤醒ATECC608A并使用Slot 0签名", "process"),
                    ("ATECC608A进入Sleep模式", "process"),
                    ("将原始(r,s)签名编码为DER/ASN.1", "process"),
                    ("构建signature_envelope", "process"),
                    ("发送到signedQueue并继续循环", "queue"),
                ],
            ),
        ),
        ("代码5-3 CommTask上传流程伪代码", "流程图5-3 CommTask上传流程图", "flow_5_3_commtask.png", comm_chart),
        (
            "代码5-4 Python端规范化哈希核心实现",
            "流程图5-4 Python端规范化哈希流程图",
            "flow_5_4_python_hash.png",
            lambda p: canonical_branch_chart(p, "输入Python字典数据", "json.dumps后计算SHA-256"),
        ),
        (
            "代码5-5 TypeScript端规范化哈希核心实现",
            "流程图5-5 TypeScript端规范化哈希流程图",
            "flow_5_5_ts_hash.png",
            lambda p: canonical_branch_chart(p, "输入TypeScript对象数据", "TextEncoder编码后调用Web Crypto SHA-256"),
        ),
        ("代码5-6 canonical_hash冲突处理逻辑", "流程图5-6 canonical_hash冲突处理流程图", "flow_5_6_conflict.png", conflict_chart),
        ("代码5-7 ECDSA签名验证核心实现", "流程图5-7 ECDSA签名验证流程图", "flow_5_7_ecdsa.png", ecdsa_chart),
        ("代码5-8 HMAC验证核心实现", "流程图5-8 HMAC验证流程图", "flow_5_8_hmac.png", hmac_chart),
        ("代码5-9 锚定任务核心流程", "流程图5-9 锚定任务核心流程图", "flow_5_9_anchor.png", anchor_chart),
    ]

    for original_caption, new_caption, filename, maker in specs:
        img_path = OUT_DIR / filename
        maker(img_path)
        images[original_caption] = img_path
    return images


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_caption(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_picture(paragraph, img_path: Path) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(img_path), width=Cm(14.2))


def next_paragraph(paragraphs, idx: int):
    return paragraphs[idx + 1] if idx + 1 < len(paragraphs) else None


def replace_in_doc(images: dict[str, Path]) -> None:
    doc = Document(INPUT_DOCX)
    replacements = {
        "代码5-1 SensorTask任务伪代码": "流程图5-1 SensorTask任务流程图",
        "代码5-2 SignTask任务伪代码": "流程图5-2 SignTask任务流程图",
        "代码5-3 CommTask上传流程伪代码": "流程图5-3 CommTask上传流程图",
        "代码5-4 Python端规范化哈希核心实现": "流程图5-4 Python端规范化哈希流程图",
        "代码5-5 TypeScript端规范化哈希核心实现": "流程图5-5 TypeScript端规范化哈希流程图",
        "代码5-6 canonical_hash冲突处理逻辑": "流程图5-6 canonical_hash冲突处理流程图",
        "代码5-7 ECDSA签名验证核心实现": "流程图5-7 ECDSA签名验证流程图",
        "代码5-8 HMAC验证核心实现": "流程图5-8 HMAC验证流程图",
        "代码5-9 锚定任务核心流程": "流程图5-9 锚定任务核心流程图",
    }
    paragraphs = doc.paragraphs
    seen: set[str] = set()

    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "伪代码流程如下：":
            clear_paragraph(paragraph)
            run = paragraph.add_run("任务流程图如下：")
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)
            continue
        if text not in replacements:
            continue
        following = next_paragraph(paragraphs, i)
        if following is None:
            raise RuntimeError(f"未找到 {text} 后的代码段落")
        set_picture(paragraph, images[text])
        set_caption(following, replacements[text])
        seen.add(text)

    missing = sorted(set(replacements) - seen)
    if missing:
        raise RuntimeError("未找到以下代码标题：" + "、".join(missing))

    doc.save(OUTPUT_DOCX)


def main() -> None:
    images = generate_images()
    replace_in_doc(images)
    print(f"saved: {OUTPUT_DOCX}")
    print(f"images: {OUT_DIR}")


if __name__ == "__main__":
    main()
